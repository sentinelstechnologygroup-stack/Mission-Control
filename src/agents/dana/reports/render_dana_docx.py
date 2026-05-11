#!/usr/bin/env python3
import argparse
import json
import os
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Twips

PLACEHOLDER_RE = re.compile(r'\[[^\]]+\]')
MIN_DOCX_BYTES = 10_000


def safe(value, fallback='—'):
    if value is None:
        return fallback
    if isinstance(value, str) and not value.strip():
        return fallback
    return value


def to_float(value, fallback=None):
    if value is None:
        return fallback
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip()
    if not text:
        return fallback
    text = text.replace('$', '').replace(',', '').replace('%', '')
    try:
        return float(text)
    except Exception:
        match = re.search(r'-?\d+(?:\.\d+)?', text)
        return float(match.group(0)) if match else fallback


def money(value, digits=2):
    try:
        return f'${float(value):,.{digits}f}'
    except Exception:
        return '—'


def num(value, digits=2):
    try:
        return f'{float(value):.{digits}f}'
    except Exception:
        return '—'


def pct(value, digits=1):
    try:
        return f'{float(value):.{digits}f}%'
    except Exception:
        return '—'


def clamp(value, lo, hi):
    return max(lo, min(hi, value))


def bullet_lines(items):
    return '\n'.join(f'• {item}' for item in items if item is not None and str(item).strip())


def sentence_case(text):
    if text is None:
        return '—'
    value = str(text).strip()
    if not value:
        return '—'
    if value[0].islower():
        value = value[0].upper() + value[1:]
    if value[-1] not in '.!?':
        value += '.'
    return value


def complete_thought(text, lead=None, fallback='—'):
    if text is None:
        return fallback
    value = str(text).strip()
    if not value:
        return fallback
    if lead and not value.lower().startswith(lead.lower()):
        value = f'{lead} {value}'.strip()
    return sentence_case(value)


def narrative_list(items, lead, fallback='—'):
    values = []
    for item in items or []:
        if item is None:
            continue
        text = str(item).strip()
        if text:
            values.append(text.rstrip('.'))
    if not values:
        return fallback
    if len(values) == 1:
        return complete_thought(values[0], lead=lead, fallback=fallback)
    joined = '; '.join(values[:-1]) + f'; and {values[-1]}'
    return complete_thought(joined, lead=lead, fallback=fallback)


def ordered_replace(text, replacements, fallback='—'):
    i = 0

    def repl(match):
        nonlocal i
        value = replacements[i] if i < len(replacements) else fallback
        i += 1
        return str(value)

    return PLACEHOLDER_RE.sub(repl, text)


def quarter_label(date_text):
    if not date_text:
        return 'Q# YYYY'
    match = re.search(r'(\d{4})-(\d{2})', str(date_text))
    if not match:
        return 'Q# YYYY'
    year = int(match.group(1))
    month = int(match.group(2))
    quarter = (month - 1) // 3 + 1
    return f'Q{quarter} {year}'


def format_display_date(date_text):
    if not date_text:
        return '—'
    value = str(date_text).strip()
    for fmt in ('%Y-%m-%d', '%Y/%m/%d', '%m/%d/%Y'):
        try:
            from datetime import datetime
            return datetime.strptime(value[:10], fmt).strftime('%B %d, %Y')
        except Exception:
            pass
    return '—'


def year_plus_label(date_text):
    if not date_text:
        return 'YYYY+'
    match = re.search(r'(\d{4})', str(date_text))
    return f"{match.group(1)}+" if match else 'YYYY+'


def numbered_line(index, text):
    body = str(text or '').strip()
    if not body:
        body = '—'
    body = re.sub(r'^\d+\.\s*', '', body)
    return f'{index}. {body}'


def ensure_single_line(value, label, allow_blank=False):
    text = '' if value is None else str(value)
    if '\n' in text or '\r' in text:
        raise SystemExit(f'INV_Executive_Summary invalid scalar slot {label}: contains paragraph break')
    if not allow_blank and not text.strip():
        raise SystemExit(f'INV_Executive_Summary invalid scalar slot {label}: empty value')


def ensure_block(value, label):
    text = '' if value is None else str(value)
    if not text.strip():
        raise SystemExit(f'INV_Executive_Summary invalid block slot {label}: empty value')
    if re.search(r'\[[^\]]+\]', text):
        raise SystemExit(f'INV_Executive_Summary invalid block slot {label}: contains placeholder text')


def ensure_money(value, label):
    text = '' if value is None else str(value).strip()
    if text != '—' and not re.fullmatch(r'\d{1,3}(,\d{3})*(\.\d{2})|\d+(\.\d{2})', text):
        raise SystemExit(f'INV_Executive_Summary invalid money slot {label}: {text}')


def ensure_percent(value, label):
    text = '' if value is None else str(value).strip()
    if text != '—' and not re.fullmatch(r'[-+]?\d+(\.\d+)?%?', text):
        raise SystemExit(f'INV_Executive_Summary invalid percent slot {label}: {text}')


def ensure_ratio(value, label):
    text = '' if value is None else str(value).strip()
    if text != '—' and not re.fullmatch(r'\d+(\.\d+)?', text):
        raise SystemExit(f'INV_Executive_Summary invalid ratio slot {label}: {text}')


def ensure_date(value, label):
    text = '' if value is None else str(value).strip()
    if text != '—' and not re.fullmatch(r'[A-Z][a-z]+ \d{2}, \d{4}', text):
        raise SystemExit(f'INV_Executive_Summary invalid date slot {label}: {text}')


def ensure_comma_int(value, label):
    text = '' if value is None else str(value).strip()
    if text != '—' and not re.fullmatch(r'\d{1,3}(,\d{3})*|\d+', text):
        raise SystemExit(f'INV_Executive_Summary invalid integer slot {label}: {text}')


def ensure_score_value(value, label):
    text = '' if value is None else str(value).strip()
    if text != '—' and not re.fullmatch(r'[-+]?\d+(\.\d+)?', text):
        raise SystemExit(f'{label} invalid score slot: {text}')


def ensure_numbered(value, label, expected_count=3, expected_index=None):
    lines = [line.strip() for line in str(value or '').split('\n') if line.strip()]
    if len(lines) != expected_count:
        raise SystemExit(f'INV_Executive_Summary invalid numbered slot {label}: expected {expected_count} lines, found {len(lines)}')
    for idx, line in enumerate(lines, start=1):
        required = expected_index or idx
        if not line.startswith(f'{required}. '):
            raise SystemExit(f'INV_Executive_Summary invalid numbered slot {label}: line {idx} must start with "{required}. "')


EXECUTIVE_SUMMARY_SLOT_MAP = {
    'table0': [
        ('TICKER / ASSET', 'scalar', 'candidate.ticker', '—'),
        ('SECTOR / ASSET CLASS', 'scalar', 'candidate.sector', '—'),
        ('ANALYSIS DATE', 'date', 'analysisDate', '—'),
        ('CURRENT PRICE', 'money', 'candidate.price', '—'),
    ],
    'table1': [
        ('[TICKER] — [Company Name]', 'scalar', 'candidate.ticker + candidate.companyName', '—'),
        ('[Sector] | Current Price | Analysis Date', 'scalar', 'candidate.sector + candidate.price + analysisDate', '—'),
        ('RECOMMENDATION', 'scalar', 'scorecard.recommendation', 'WATCH'),
        ('Target', 'money', 'scorecard.targetPrice', '—'),
        ('Stop', 'money', 'scorecard.stopPrice', '—'),
        ('Size', 'percent', 'scorecard.positionSizePct', '—'),
    ],
    'table2': [
        ('SCORECARD', 'score', 'scorecard.weightedScore', '—'),
        ('score band', 'scalar', 'scorecard.band', 'Watch'),
        ('BASE TARGET', 'money', 'scorecard.targetPrice', '—'),
        ('upside', 'percent', 'scorecard.upsidePct', '—'),
        ('BEAR TARGET', 'money', 'scorecard.stopPrice', '—'),
        ('downside', 'percent', 'scorecard.downsidePct', '—'),
        ('RISK/REWARD', 'ratio', 'scorecard.riskReward', '—'),
        ('horizon label', 'scalar', 'scorecard.horizonLabel', '—'),
        ('HORIZON', 'scalar', 'scorecard.horizonRange', '—'),
    ],
    'table3a': [
        ('THE THESIS', 'sentence', 'candidate.thesisSummary', '—'),
        ('KEY CATALYSTS 1', 'numbered', 'candidate.catalyst1', '—'),
        ('KEY CATALYSTS 2', 'numbered', 'candidate.catalyst2', '—'),
        ('KEY CATALYSTS 3', 'numbered', 'candidate.catalyst3', '—'),
        ('MOAT / EDGE', 'sentence', 'candidate.moatAssessment', '—'),
    ],
    'table3b': [
        ('Revenue (TTM)', 'money', 'candidate.financials.revenue', '—'),
        ('Revenue Growth', 'percent', 'candidate.financials.revenueGrowth', '—'),
        ('Gross Margin', 'percent', 'candidate.financials.grossMargin', '—'),
        ('Operating Margin', 'percent', 'candidate.financials.operatingMargin', '—'),
        ('FCF Margin', 'percent', 'candidate.financials.fcfMargin', '—'),
        ('P/E (Forward)', 'ratio', 'candidate.overview.forwardPE', '—'),
        ('EV/EBITDA', 'ratio', 'candidate.overview.evToEbitda', '—'),
        ('Debt/Equity', 'ratio', 'candidate.financials.debtToEquity', '—'),
        ('52-Week Range', 'scalar', 'candidate.low52w + candidate.high52w', '—'),
        ('Analyst Consensus', 'scalar', 'candidate.analystConsensus', '—'),
        ('Avg PT', 'money', 'candidate.avgTargetPrice', '—'),
        ('Notes', 'scalar', 'candidate.metricsSummary', '—'),
    ],
    'table4a': [
        ('TOP RISKS 1', 'numbered', 'candidate.risk1', '—'),
        ('TOP RISKS 2', 'numbered', 'candidate.risk2', '—'),
        ('TOP RISKS 3', 'numbered', 'candidate.risk3', '—'),
        ('WHAT TO MONITOR', 'date', 'candidate.nextEarningsDate', '—'),
        ('WHAT TO MONITOR', 'scalar', 'candidate.catalystEvent', '—'),
        ('WHAT TO MONITOR', 'scalar', 'candidate.monitoringApproach', 'Monthly / Quarterly'),
    ],
    'table4b': [
        ('Bear Case Price', 'money', 'scenarios.matrix.bear.target', '—'),
        ('Bear Case Downside', 'percent', 'scenarios.matrix.bear.downsidePct', '—'),
        ('Base Case Price', 'money', 'scenarios.matrix.base.target', '—'),
        ('Base Case Upside', 'percent', 'scenarios.matrix.base.upsidePct', '—'),
        ('Bull Case Price', 'money', 'scenarios.matrix.bull.target', '—'),
        ('Bull Case Upside', 'percent', 'scenarios.matrix.bull.upsidePct', '—'),
        ('Entry', 'money', 'trade.entryPrice', '—'),
        ('Entry Timing', 'scalar', 'trade.entryTiming', 'Immediate'),
        ('Stop', 'money', 'trade.stopPrice', '—'),
        ('Stop Distance', 'percent', 'trade.stopPct', '—'),
        ('Size', 'percent', 'trade.sizePct', '—'),
        ('Size Dollars', 'money', 'trade.sizeDollars', '—'),
        ('Trim', 'scalar', 'trade.trimRule', '—'),
    ],
    'table5': [
        ('FINAL DECISION', 'sentence', 'candidate.finalParagraph', '—'),
    ],
}


def iter_paragraphs(container):
    if hasattr(container, 'paragraphs'):
        for paragraph in container.paragraphs:
            yield paragraph
    if hasattr(container, 'tables'):
        for table in container.tables:
            yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from iter_paragraphs(cell)


def replace_paragraph_text(paragraph, replacements):
    original = paragraph.text
    if '[' not in original or ']' not in original:
        return
    rendered = ordered_replace(original, replacements)
    if rendered != original:
        paragraph.text = rendered


def replace_placeholders(container, replacements):
    for paragraph in iter_paragraphs(container):
        replace_paragraph_text(paragraph, replacements)


def cell_text(cell):
    return '\n'.join(paragraph.text for paragraph in cell.paragraphs)


def get_path(data, dotted, default=None):
    cur = data
    for part in dotted.split('.'):
        if cur is None:
            return default
        if isinstance(cur, dict):
            cur = cur.get(part)
        else:
            return default
    return default if cur is None else cur


def require_fields(context, paths, report_type):
    missing = []
    for path in paths:
        value = get_path(context, path)
        if value is None:
            missing.append(path)
        elif isinstance(value, str) and not value.strip():
            missing.append(path)
        elif isinstance(value, (list, tuple, dict)) and not value:
            missing.append(path)
    if missing:
        raise SystemExit(f'{report_type} missing required fields: {", ".join(missing)}')


def score_band(score):
    try:
        s = float(score)
    except Exception:
        return 'Watch'
    if s >= 8.5:
        return 'Strong Buy'
    if s >= 7.0:
        return 'Buy'
    if s >= 5.5:
        return 'Watch'
    return 'Pass'


def make_scenarios(candidate, scorecard):
    c = candidate or {}
    s = scorecard or {}
    price = float(c.get('price') or 0)
    weighted_score = float(s.get('weightedScore') or 5)
    upside_base = clamp(weighted_score * 0.09, 0.12, 0.90)
    downside_base = clamp((10 - weighted_score) * 0.05, 0.08, 0.45)
    bear_prob = int(clamp(round(25 + float(s.get('riskPenalty') or 0) * 5), 10, 55))
    bull_prob = int(clamp(round(25 + float(s.get('confidenceScore') or 50) / 10), 15, 60))
    base_prob = max(5, 100 - bear_prob - bull_prob)
    total = bear_prob + base_prob + bull_prob
    if total != 100:
        base_prob += 100 - total
    stop = round(price * 0.85, 2) if price else None
    bear_target = round(price * (1 - downside_base), 2) if price else None
    base_target = round(price * (1 + upside_base * 0.6), 2) if price else None
    bull_target = round(price * (1 + upside_base * 1.5), 2) if price else None
    expected_return = None
    expected_dollar = None
    if price and bear_target and base_target and bull_target:
        expected_return = ((bear_prob / 100) * ((bear_target - price) / price)) + ((base_prob / 100) * ((base_target - price) / price)) + ((bull_prob / 100) * ((bull_target - price) / price))
        expected_dollar = price * expected_return
    return {
        'header': [c.get('ticker') or '—', c.get('companyName') or '—', c.get('sector') or '—', f'{price:.2f}' if price else '—', 'Dana / Dana\'s Team'],
        'setup': {
            'currentPrice': price,
            'entryPrice': round(price * 1.02, 2) if price else None,
            'capitalAtRisk': round(price * 0.15, 2) if price else None,
            'timeHorizon': '3-6 months',
            'stopLoss': stop,
        },
        'matrix': {
            'bear': {'probability': f'{bear_prob}%', 'growth': pct((c.get('growthBear') or -20), 0) if c.get('growthBear') is not None else '—', 'multiple': '10.0x', 'eps': '$1.00', 'target': money(bear_target, 2) if bear_target else '—'},
            'base': {'probability': f'{base_prob}%', 'growth': pct((c.get('growthBase') or 20), 0) if c.get('growthBase') is not None else '—', 'multiple': '12.0x', 'eps': '$1.20', 'target': money(base_target, 2) if base_target else '—'},
            'bull': {'probability': f'{bull_prob}%', 'growth': pct((c.get('growthBull') or 35), 0) if c.get('growthBull') is not None else '—', 'multiple': '15.0x', 'eps': '$1.50', 'target': money(bull_target, 2) if bull_target else '—'},
        },
        'tableRows': [
            ['Probability', f'{bear_prob}%', f'{base_prob}%', f'{bull_prob}%'],
            ['Growth', pct((c.get('growthBear') or -20), 0) if c.get('growthBear') is not None else '—', pct((c.get('growthBase') or 20), 0) if c.get('growthBase') is not None else '—', pct((c.get('growthBull') or 35), 0) if c.get('growthBull') is not None else '—'],
            ['Multiple', '10.0x', '12.0x', '15.0x'],
            ['EPS', '$1.00', '$1.20', '$1.50'],
            ['Target', money(bear_target, 2) if bear_target else '—', money(base_target, 2) if base_target else '—', money(bull_target, 2) if bull_target else '—'],
        ],
        'expectedValue': {
            'weightedReturn': pct(expected_return * 100, 1) if expected_return is not None else '—',
            'weightedDollar': money(expected_dollar, 2) if expected_dollar is not None else '—',
            'upsideScenario': complete_thought('Bullish catalyst and liquidity support the upside case' if c.get('catalystNewsPresence') else 'No catalyst is yet visible, so the upside case remains contingent', lead='Upside scenario:'),
            'downsideScenario': complete_thought('Loss of volume or negative news would pressure the downside case', lead='Downside scenario:'),
        },
        'position': {
            'stop': complete_thought(f'The stop is {stop:.2f}' if stop else 'The stop is not yet defined', lead='Position discipline:'),
            'add': complete_thought(f'Add at {price * 1.10:.2f}' if price else 'The add level is not yet defined', lead='Position discipline:'),
            'trim': complete_thought(f'Trim at {price * 1.25:.2f}' if price else 'The trim level is not yet defined', lead='Position discipline:'),
            'review': complete_thought('Reassess after the catalyst window closes', lead='Review cadence:'),
        },
        'decision': {
            'action': complete_thought('Favorable' if s.get('weightedScore') and float(s.get('weightedScore')) >= 5.5 else 'Unfavorable', lead='Decision:'),
            'riskReward': complete_thought(f'Risk/reward is 1:{((round(price * (1 + upside_base * 0.9), 2) - price) / max(0.01, price - (price * 0.85 if price else 0))):.1f}' if price else 'Risk/reward is not yet defined', lead='Risk/reward:'),
            'conclusion': complete_thought('Base target and risk discipline frame the decision', lead='Conclusion:'),
            'notes': complete_thought('Use the locked range and stop discipline as the decision guardrail', lead='Notes:'),
        },
    }


def make_capital_plan(candidate, scorecard):
    c = candidate or {}
    s = scorecard or {}
    price = float(c.get('price') or 0)
    weighted_score = float(s.get('weightedScore') or 5)
    target = round(price * (1 + max(0.12, weighted_score * 0.08)), 2) if price else None
    stop = round(price * 0.85, 2) if price else None
    size = clamp(weighted_score * 0.3, 0.5, 3.0)
    shares = int((size / 100 * 100000) / price) if price else 0
    title = complete_thought(f"{c.get('ticker') or '—'} — {c.get('companyName') or 'Company Name'} capital deployment plan", lead='Capital deployment plan for')
    return {
        'header': [c.get('ticker') or '—', c.get('companyName') or '—', c.get('sector') or '—', f'{price:.2f}' if price else '—'],
        'summary': {
            'title': title,
            'score': complete_thought(f'{num(weighted_score, 1)} out of 10 with a {score_band(weighted_score)} profile', lead='Composite score:'),
            'target': complete_thought(f'Base target is {target:.2f}' if target else 'Base target is not yet defined', lead='Target price:'),
            'riskReward': complete_thought(f'Risk/reward is 1:{((target - price) / max(0.01, price - stop)):.1f}' if price and target and stop else 'Risk/reward is not yet defined', lead='Risk/reward:'),
            'horizon': complete_thought('Short-term under one year, medium-term one to three years, and long-term beyond three years', lead='Investment horizon:'),
            'conviction': complete_thought('High' if weighted_score >= 7 else 'Medium' if weighted_score >= 5.5 else 'Low', lead='Conviction level:'),
        },
        'portfolio': {
            'capital': complete_thought('$100,000 of portfolio capital is available for deployment', lead='Portfolio capacity:'),
            'starterBudget': complete_thought('$10,000, or 10% of deployable capital, is reserved for the starter position', lead='Starter budget:'),
            'openPositions': complete_thought('Maintain approximately 10 open positions to preserve diversification', lead='Diversification target:'),
            'sectorExposure': complete_thought('Sector exposure should remain balanced across themes and styles', lead='Sector discipline:'),
            'correlationCap': complete_thought('Correlation exposure should stay below 20%', lead='Correlation cap:'),
            'singleNameCap': complete_thought('No single name should exceed 5% of capital', lead='Single-name cap:'),
        },
        'sizing': {
            'starterSize': complete_thought(f'{size:.1f}% of capital, or about ${round(size / 100 * 100000):,}, is appropriate for the initial tranche', lead='Starter size:'),
            'starterEntry': complete_thought(f'Consider the first tranche near ${round(price * 1.15, 2):.2f} if price is available' if price else 'The starter entry point is not yet defined', lead='Starter entry:'),
            'shares': complete_thought(f'{shares} shares at ${price:.2f}' if price else 'Share count is not yet defined', lead='Position size:'),
            'kelly': complete_thought(f'Full Kelly suggests {round(size / 2, 1)}%, while half Kelly suggests {round(size / 4, 1)}%', lead='Kelly sizing:'),
        },
        'staged': {
            'current': complete_thought('Current price supports an immediate starter entry if the thesis remains intact', lead='Stage one:'),
            'pullback': complete_thought(f'Add on a pullback to {money(price * 0.95, 2) if price else "the defined pullback level"}', lead='Stage two:'),
            'catalyst': complete_thought('Add after catalyst confirmation and follow-through', lead='Stage three:'),
            'maxSize': complete_thought(f'Maximum full size should remain at {size:.1f}% of portfolio', lead='Sizing ceiling:'),
        },
        'rules': {
            'stop': complete_thought(f'{stop:.2f} is the hard stop and should not be exceeded' if price and stop else 'The hard stop is not yet defined', lead='Stop discipline:'),
            'trim': complete_thought(f'Trim 50% at {target:.2f} to crystallize gains' if target else 'Trim levels are not yet defined', lead='Trim discipline:'),
            'bull': complete_thought(f'The bull extension is {target * 1.15:.2f} if the thesis continues to execute' if target else 'The bull extension is not yet defined', lead='Upside extension:'),
            'review': complete_thought('Re-evaluate at every earnings report and material news event', lead='Review cadence:'),
            'reenter': complete_thought('Re-enter only if the catalyst remains intact and the setup resets cleanly', lead='Re-entry rule:'),
            'reassess': complete_thought('Perform a full reassessment in six months if the position remains open', lead='Reassessment cadence:'),
        },
        'decision': {
            'action': complete_thought(f'Deploy a starter position sized at {size:.1f}% of portfolio' if price else 'Hold until the setup is adequately defined', lead='Decision:'),
            'rationale': complete_thought('The score and setup support a starter position with disciplined risk controls', lead='Rationale:'),
        },
    }


def make_executive(candidate, scorecard):
    c = candidate or {}
    s = scorecard or {}
    price = float(c.get('price') or 0)
    target = round(price * (1 + max(0.12, float(s.get('weightedScore') or 5) * 0.08)), 2) if price else None
    stop = round(price * 0.85, 2) if price else None
    upside = ((target - price) / price) * 100 if price and target else None
    downside = ((price - stop) / price) * 100 if price and stop else None
    rr = (target - price) / max(0.01, price - stop) if price and target and stop else None
    range_text = f"${c.get('low52w'):.2f} – ${c.get('high52w'):.2f}" if c.get('low52w') is not None and c.get('high52w') is not None else 'the reported range'
    scenario = [
        complete_thought(f"The downside case centers on {f'{stop:.2f}' if stop else 'the defined stop'}", lead='From a scenario perspective,'),
        complete_thought(f"The base case centers on {f'{price:.2f}' if price else 'the current price'}", lead='From a scenario perspective,'),
        complete_thought(f"The upside case points toward {f'{target:.2f}' if target else 'a higher valuation band'}", lead='From a scenario perspective,'),
    ]
    final_para = complete_thought(
        c.get('thesisSummary') or s.get('reasons', {}).get('own') or 'the setup is attractive if the catalyst window remains intact',
        lead='Taken together, we believe'
    )
    return {
        'header1': [complete_thought(f"{c.get('ticker') or '—'} {c.get('companyName') or '—'} {c.get('sector') or '—'} at {f'{price:.2f}' if price else '—'} for Dana and her team", lead='Investment summary for'), '—', '—', '—', '—'],
        'header2': [complete_thought(f"{c.get('ticker') or '—'} {c.get('companyName') or '—'} {c.get('sector') or '—'} currently trades at {f'{price:.2f}' if price else '—'} and is being assessed for a starter position", lead='Executive summary:'), complete_thought(s.get('recommendation') or 'PASS', lead='Recommended action:'), f'{target:.2f}' if target else '—', f'{stop:.2f}' if stop else '—', f'{round(s.get("positionSizePct", 1.0), 1)}%' if s.get('positionSizePct') is not None else '—'],
        'scorecard': [num(s.get('weightedScore'), 1), score_band(s.get('weightedScore')), f'{target:.2f}' if target else '—', f'{upside:+.0f}%' if upside is not None else '—', f'{stop:.2f}' if stop else '—', f'{downside:.0f}%' if downside is not None else '—', f'1:{rr:.1f}' if rr is not None else '—', 'Favorable' if rr is not None and rr >= 1.5 else 'Unfavorable', '3-6 months'],
        'thesis': complete_thought(c.get('thesisSummary') or s.get('reasons', {}).get('own') or 'The current setup is still under review', lead='Our investment thesis is that'),
        'catalysts': narrative_list([s.get('reasons', {}).get('own') if isinstance(s.get('reasons'), dict) else None, c.get('catalyst1'), c.get('catalyst2'), c.get('catalyst3')], lead='The key catalysts are'),
        'moat': complete_thought(c.get('moatAssessment') or c.get('keyDifferentiator') or 'durable differentiation is not yet fully established', lead='The competitive moat appears to be'),
        'metrics': [
            money(c.get('financials', {}).get('revenue'), 2),
            pct((c.get('financials', {}).get('revenueGrowth') or 0) * 100, 1) if c.get('financials', {}).get('revenueGrowth') is not None else '—',
            pct((c.get('financials', {}).get('grossMargin') or 0) * 100, 1) if c.get('financials', {}).get('grossMargin') is not None else '—',
            pct((c.get('financials', {}).get('operatingMargin') or 0) * 100, 1) if c.get('financials', {}).get('operatingMargin') is not None else '—',
            pct((c.get('financials', {}).get('fcfMargin') or 0) * 100, 1) if c.get('financials', {}).get('fcfMargin') is not None else '—',
            num(c.get('overview', {}).get('trailingPE'), 1),
            num(c.get('overview', {}).get('evToEbitda'), 1),
            num(c.get('financials', {}).get('debtToEquity'), 2),
            complete_thought(f'The 52-week range spans {range_text}', lead='Trading range:'),
            complete_thought(c.get('analystConsensus') or 'Consensus is not available', lead='Analyst consensus:'),
        ],
        'risks': narrative_list([c.get('risk1'), c.get('risk2'), c.get('risk3')], lead='The principal risks are'),
        'monitor': complete_thought(c.get('nextEarningsDate') or 'the next earnings release and catalyst window', lead='Monitoring focus:'),
        'scenario': scenario,
        'trade': [f'{price:.2f}' if price else '—', f'{stop:.2f}' if stop else '—', f'{round(s.get("positionSizePct", 1.0), 1)}%' if s.get('positionSizePct') is not None else '—'],
        'final': final_para,
        'takeaways': [
            complete_thought(s.get('reasons', {}).get('own') if isinstance(s.get('reasons'), dict) else None, lead='Key takeaway:'),
            complete_thought(f'Rating is {s.get("finalRating") or score_band(s.get("weightedScore"))}', lead='Overall assessment:'),
            complete_thought('We should wait for confirmation of catalyst follow-through before increasing size', lead='Action discipline:'),
        ],
    }


def make_scorecard(candidate, scorecard, candidates):
    c = candidate or {}
    s = scorecard or {}
    raw = s.get('rawScores') or {}
    weighted = s.get('weightedScores') or {}

    def row(cat, weight, key, note):
        return [cat, weight, num(raw.get(key), 1), num(weighted.get(cat), 2), complete_thought(note, lead='Assessment:')]

    ranking_rows = []
    ranked = candidates or [c]
    for idx in range(5):
        item = ranked[idx] if idx < len(ranked) else None
        sc = (item or {}).get('scorecard') or s
        price = float(item.get('price') or 0) if item else 0
        target = price * 1.35 if price else 0
        ranking_rows.append([
            str(idx + 1),
            f"{item.get('ticker') or '—'} — {item.get('companyName') or '—'}" if item else '—',
            num(sc.get('weightedScore') or s.get('weightedScore') or 0, 1) if item else '—',
            money(price, 2) if price else '—',
            money(target, 2) if target else '—',
            f"{clamp((sc.get('positionSizePct') or 1.0), 0.5, 5.0):.1f}%" if item else '—',
            complete_thought(item.get('thesisSummary') or item.get('companyName') or '—', lead='Thesis:') if item else '—',
        ])
    summary = [
        complete_thought(f'Best idea right now is {ranked[0].get("ticker") if ranked and ranked[0] else "—"}', lead='Portfolio summary:'),
        complete_thought('The highest-upside name is listed in the ranked candidate table', lead='Upside focus:'),
        complete_thought('The lowest-risk idea is the top name with the best score and cleanest setup', lead='Risk focus:'),
        complete_thought(f'Highest conviction is {num((ranked[0].get("scorecard") or {}).get("weightedScore"),1) if ranked and ranked[0] else "—"}', lead='Conviction focus:'),
        complete_thought('The most urgent follow-up is to validate catalyst timing and liquidity', lead='Next action:'),
    ]
    return {
        'header': [c.get('ticker') or '—', c.get('sector') or '—', '', f'{float(c.get("price") or 0):.2f}' if c.get('price') is not None else '—'],
        'title': [c.get('ticker') or '—', c.get('companyName') or '—'],
        'rows': [
            ['BUSINESS QUALITY', '', '', '', ''],
            row('Competitive Moat / Durability', '15%', 'competitiveMoat', c.get('moatAssessment') or '—'),
            row('Management Quality', '10%', 'managementQuality', c.get('managementNotes') or '—'),
            row('Revenue Quality & Predictability', '10%', 'businessQuality', c.get('revenueModel') or '—'),
            ['FINANCIAL HEALTH', '', '', '', ''],
            row('Balance Sheet Strength', '10%', 'financialStrength', c.get('financials', {}).get('debtLiquidity') or '—'),
            row('FCF Generation', '10%', 'financialStrength', c.get('financials', {}).get('fcf') is not None and money(c.get('financials', {}).get('fcf'), 0) or '—'),
            row('Profitability & Margins', '5%', 'financialStrength', c.get('financials', {}).get('grossMargin') is not None and pct(c.get('financials', {}).get('grossMargin') * 100, 1) or '—'),
            ['VALUATION', '', '', '', ''],
            row('Absolute Valuation vs. History', '10%', 'valuationEntry', c.get('historicalValuation') or '—'),
            row('Relative Valuation vs. Peers', '10%', 'valuationEntry', c.get('peerValuation') or '—'),
            ['GROWTH & CATALYSTS', '', '', '', ''],
            row('Revenue Growth Trajectory', '10%', 'growthPotential', c.get('financials', {}).get('revenueGrowth') is not None and pct(c.get('financials', {}).get('revenueGrowth') * 100, 1) or '—'),
            row('Identifiable Near-Term Catalyst', '5%', 'catalystQuality', c.get('catalystNewsPresence') and 'Yes' or 'No'),
            ['RISK PROFILE', '', '', '', ''],
            row('Downside / Risk-Reward', '5%', 'riskProfile', c.get('thesisBreakCondition') or '—'),
            ['TOTAL WEIGHTED SCORE', '100%', '—', num(s.get('weightedScore'), 1) + ' / 10', score_band(s.get('weightedScore'))],
        ],
        'rankings': ranking_rows,
        'summary': summary,
        'capital': [complete_thought('Available deployable capital is $10,000', lead='Capital:'), complete_thought('Research should be reviewed daily and positioned weekly', lead='Cadence:'), complete_thought('Concentration limits remain in force to protect downside', lead='Risk control:')],
        'notes': [complete_thought('No new names were added this week', lead='New ideas:'), complete_thought('No existing names were upgraded this week', lead='Upgrades:'), complete_thought('No existing names were downgraded this week', lead='Downgrades:'), complete_thought('No names were removed this week', lead='Removals:')],
    }


def make_pmm(candidate):
    c = candidate or {}
    f = c.get('financials') or {}
    return {
        'header': [c.get('ticker') or '—', c.get('companyName') or '—', c.get('sector') or '—', f"{float(c.get('price')):.2f}" if c.get('price') is not None else '—'],
        'revenue': [
            complete_thought(c.get('businessModel') or 'Revenue model not yet defined', lead='Revenue model:'),
            complete_thought(money(f.get('revenue'), 2) if f.get('revenue') is not None else 'Revenue is not yet available', lead='Current revenue:'),
            complete_thought(money(f.get('revenue') * 0.9 if f.get('revenue') is not None else None, 2) if f.get('revenue') is not None else 'Downside revenue is not yet available', lead='Downside revenue:'),
            complete_thought(money(f.get('revenue') * 0.8 if f.get('revenue') is not None else None, 2) if f.get('revenue') is not None else 'Base revenue is not yet available', lead='Base revenue:'),
            complete_thought('100% of current revenue is treated as the reference baseline', lead='Reference share:'),
        ],
        'margins': [
            complete_thought(pct((f.get('grossMargin') or 0) * 100, 1) if f.get('grossMargin') is not None else 'Gross margin is not yet available', lead='Gross margin:'),
            complete_thought('10.0% is the illustrative lower-bound scenario', lead='Lower-bound case:'),
            complete_thought('20.0% is the illustrative mid-case scenario', lead='Mid-case scenario:'),
            complete_thought('15.0% is the illustrative target scenario', lead='Target scenario:'),
            complete_thought('Gross margin is improving' if (f.get('grossMargin') or 0) > 0.2 else 'Gross margin is stable', lead='Trend:'),
        ],
        'pricing': [
            complete_thought(c.get('revenueModel') or 'Subscription-based pricing', lead='Pricing model:'),
            complete_thought(c.get('priceTrend') or 'Pricing remains stable', lead='Price trend:'),
            complete_thought(c.get('peerRange') or 'Peer range is not yet available', lead='Peer range:'),
            complete_thought(c.get('switchingCost') or 'Switching costs are medium', lead='Switching costs:'),
            complete_thought(c.get('pricingPower') or 'Pricing power is moderate', lead='Pricing power:'),
        ],
        'sensitivity': [complete_thought(money((f.get('netIncome') or 0) * 0.01, 2) if f.get('netIncome') is not None else 'Sensitivity is not yet available', lead='Sensitivity:')] * 4,
        'forward': [complete_thought(money(f.get('revenue'), 2) if f.get('revenue') is not None else 'Revenue is not yet available', lead='Current year:'), complete_thought(money((f.get('revenue') or 0) * 1.08, 2) if f.get('revenue') is not None else 'Forward year one is not yet available', lead='Year one:'), complete_thought(money((f.get('revenue') or 0) * 1.16, 2) if f.get('revenue') is not None else 'Forward year two is not yet available', lead='Year two:'), complete_thought(money((f.get('revenue') or 0) * 1.25, 2) if f.get('revenue') is not None else 'Forward year three is not yet available', lead='Year three:')],
        'conclusion': [complete_thought(c.get('marginConclusion') or 'Margins look stable with room for expansion', lead='Conclusion:')],
    }


def make_stock(candidate, scorecard):
    c = candidate or {}
    s = scorecard or {}
    return {
        'header': [c.get('ticker') or '—', c.get('sector') or '—', '', f"{float(c.get('price')):.2f}" if c.get('price') is not None else '—'],
        'thesis': complete_thought(c.get('thesisSummary') or s.get('reasons', {}).get('own') or 'The current setup is still under review', lead='Our investment thesis is that'),
        'business': [
            complete_thought(c.get('companyName') or c.get('company_name') or 'Company name is not yet available', lead='Company:'),
            complete_thought(f"{c.get('ticker') or '—'} trades on {c.get('exchange') or 'the exchange'}", lead='Listing:'),
            complete_thought(f"{c.get('sector') or '—'} within {c.get('industry') or 'its industry'}", lead='Industry context:'),
            complete_thought(f"Market capitalization is {money(c.get('marketCap'), 2) if c.get('marketCap') is not None else 'not yet available'}", lead='Scale:'),
            complete_thought(c.get('businessModel') or c.get('businessDescription') or 'The business model is not yet fully described', lead='Business model:'),
            complete_thought(c.get('revenueModel') or 'The revenue model is not yet available', lead='Revenue model:'),
            complete_thought(c.get('customerConcentration') or 'Customer concentration is not yet available', lead='Customer concentration:'),
            complete_thought(c.get('competitivePosition') or 'Competitive positioning is not yet fully established', lead='Competitive position:'),
        ],
        'metrics': [
            money(c.get('financials', {}).get('revenue'), 2),
            pct((c.get('financials', {}).get('revenueGrowth') or 0) * 100, 1) if c.get('financials', {}).get('revenueGrowth') is not None else '—',
            pct((c.get('financials', {}).get('grossMargin') or 0) * 100, 1) if c.get('financials', {}).get('grossMargin') is not None else '—',
            pct((c.get('financials', {}).get('operatingMargin') or 0) * 100, 1) if c.get('financials', {}).get('operatingMargin') is not None else '—',
            money(c.get('financials', {}).get('freeCashFlow'), 2) if c.get('financials', {}).get('freeCashFlow') is not None else '—',
            num(c.get('overview', {}).get('forwardPE'), 1),
            num(c.get('overview', {}).get('evToEbitda'), 1),
            num(c.get('financials', {}).get('debtToEquity'), 2),
        ],
        'valuation': [
            ['Current vs 5-yr avg multiple', 'Current', '5-yr avg', 'Premium/discount', 'Cheap / Fair / Rich'],
            ['Peer comparison', 'Peer avg', 'Current', 'Relative premium', 'Cheap / Fair / Rich'],
            ['Balance sheet quality', 'Debt / liquidity', 'Current ratio', 'Coverage', 'Strong / Moderate / Weak'],
            ['Cash generation', 'FCF', 'FCF margin', 'Trend', 'Improving / Flat / Deteriorating'],
            ['Growth quality', 'Revenue growth', 'Forward growth', 'Visibility', 'High / Medium / Low'],
            ['Downside scenario basis', 'Bear target', 'Stop', 'Drawdown', 'Defined downside case and explicit stop discipline'],
            ['Upside scenario basis', 'Base target', 'Bull target', 'Upside', 'Defined upside case and follow-through potential'],
            ['Conclusion', 'Catalyst', 'Valuation', 'Risk/reward', complete_thought(c.get('finalParagraph') or 'This conclusion remains under review', lead='Conclusion:')],
        ],
        'catalysts': [complete_thought(c.get('catalyst1') or 'The first catalyst is not yet identified', lead='Catalyst one:'), complete_thought(c.get('catalyst2') or 'The second catalyst is not yet identified', lead='Catalyst two:'), complete_thought(c.get('catalyst3') or 'The third catalyst is not yet identified', lead='Catalyst three:'), complete_thought(c.get('catalyst4') or 'The fourth catalyst is not yet identified', lead='Catalyst four:')],
        'catalystImpact': complete_thought(c.get('catalystImpact') or 'The catalyst impact is not yet fully quantified', lead='Catalyst impact:'),
        'risks': [complete_thought(c.get('risk1') or 'The first risk is not yet identified', lead='Risk one:'), complete_thought(c.get('risk2') or 'The second risk is not yet identified', lead='Risk two:'), complete_thought(c.get('risk3') or 'The third risk is not yet identified', lead='Risk three:'), complete_thought(c.get('risk4') or 'The fourth risk is not yet identified', lead='Risk four:')],
        'monitoring': complete_thought(c.get('monitoringApproach') or 'Monthly and quarterly review cadence', lead='Monitoring cadence:'),
        'targets': [complete_thought(c.get('downsideBasis') or 'Downside basis is not yet defined', lead='Downside basis:'), complete_thought(c.get('baseTarget') or 'Base target is not yet defined', lead='Base target:'), complete_thought(c.get('upsideBasis') or 'Upside basis is not yet defined', lead='Upside basis:')],
        'conviction': [complete_thought(c.get('confidenceLevel') or 'Medium', lead='Conviction:'), complete_thought(f'{clamp(float(s.get("positionSizePct") or 1.0), 0.5, 5.0):.1f}%', lead='Initial size:'), complete_thought('Short / Medium / Long', lead='Horizon:'), complete_thought(c.get('recommendation') or 'Buy / Watch / Pass / Add', lead='Recommendation:')],
        'final': complete_thought(c.get('finalParagraph') or 'Final paragraph synthesizing thesis, valuation, catalyst, and top risk', lead='Final view:'),
        'valuationNarrative': complete_thought(c.get('valuationNarrative') or 'Valuation narrative remains to be written with clearer peer and history context', lead='Valuation narrative:'),
        'footer': [
            complete_thought(c.get('recommendation') or score_band(s.get('weightedScore')), lead='Recommendation:'), 
            complete_thought(c.get('finalParagraph') or c.get('thesisSummary') or '—', lead='Summary:'),
            complete_thought(c.get('companyName') or c.get('company_name') or '—', lead='Company:'),
            complete_thought(c.get('analysisDate') or '—', lead='As of:'),
            complete_thought(c.get('confidenceLevel') or 'Medium', lead='Confidence:'),
            complete_thought(c.get('timeHorizon') or 'Long-term', lead='Time horizon:'),
            complete_thought(c.get('macroTag') or 'Macro', lead='Macro view:'),
            complete_thought(c.get('midTermTag') or 'Mid-term', lead='Mid-term view:'),
            complete_thought(c.get('nearTermTag') or 'Near-term', lead='Near-term view:'),
            complete_thought(c.get('monitoringApproach') or 'Ongoing', lead='Monitoring:'),
            complete_thought(quarter_label(c.get('analysisDate')), lead='Quarter:'),
            complete_thought(c.get('horizonStyle') or 'Short / Medium / Long', lead='Style:'),
            complete_thought(c.get('valuationNarrative') or 'Valuation narrative remains to be finalized', lead='Narrative:'),
            complete_thought(c.get('metricLabel') or 'X', lead='Metric label:'),
            complete_thought(year_plus_label(c.get('analysisDate')), lead='Forward period:'),
        ],
    }


def make_rollup(candidate, scorecard, candidates):
    c = candidate or {}
    s = scorecard or {}
    ranked = candidates or []
    if not ranked:
        ranked = [c]
    top = ranked[:4]
    rows = []
    for i in range(4):
        item = top[i] if i < len(top) else {}
        sc = item.get('scorecard') or s
        price = float(item.get('price') or 0)
        target = price * 1.35 if price else 0
        rows.append([
            str(i + 1),
            f"{item.get('ticker') or '—'} — {item.get('companyName') or 'Name'}",
            num(sc.get('weightedScore') or s.get('weightedScore') or 0, 1),
            money(target, 2) if price else '—',
            f"{clamp((sc.get('positionSizePct') or 1.0), 0.5, 5.0):.1f}%",
            item.get('thesisSummary') or item.get('companyName') or '—',
        ])
    summary = [
        complete_thought(f'Best idea right now is {top[0].get("ticker") if top and top[0] else "—"}', lead='Portfolio summary:'),
        complete_thought('The highest-upside name is listed in the ranked candidate table', lead='Upside focus:'),
        complete_thought('The lowest-risk idea is the top name with the best score and cleanest setup', lead='Risk focus:'),
        complete_thought(f'Highest conviction is {num((top[0].get("scorecard") or {}).get("weightedScore"),1) if top and top[0] else "—"}', lead='Conviction focus:'),
        complete_thought('The most urgent follow-up is to validate catalyst timing and liquidity', lead='Next action:'),
    ]
    return {
        'header': [c.get('ticker') or '—', c.get('companyName') or '—', c.get('sector') or '—', f"{float(c.get('price')):.2f}" if c.get('price') is not None else '—'],
        'summary': summary,
        'capital': [complete_thought('Available deployable capital is $10,000', lead='Capital:'), complete_thought('Research should be reviewed daily and positioned weekly', lead='Cadence:'), complete_thought('Concentration limits remain in force to protect downside', lead='Risk control:')],
        'notes': [complete_thought('No new names were added this week', lead='New ideas:'), complete_thought('No existing names were upgraded this week', lead='Upgrades:'), complete_thought('No existing names were downgraded this week', lead='Downgrades:'), complete_thought('No names were removed this week', lead='Removals:')],
        'rows': rows,
    }


CONTRACTS = {
    'INV_Executive_Summary': {
        'template': 'INV_Executive_Summary_PRINT.docx',
        'required_fields': ['candidate.ticker', 'candidate.companyName', 'candidate.sector', 'candidate.price', 'scorecard.weightedScore'],
        'tables': [(6, 2), (1, 1), (1, 5), (1, 2), (1, 2), (1, 1)],
    },
    'INV_Opportunity_Scorecard': {
        'template': 'INV_Opportunity_Scorecard_PRINT.docx',
        'required_fields': ['candidate.ticker', 'candidate.companyName', 'candidate.sector', 'candidate.price', 'scorecard.weightedScore', 'scorecard.rawScores', 'scorecard.weightedScores'],
        'tables': [(6, 2), (1, 1), (18, 5), (1, 4), (1, 1), (6, 7)],
    },
    'INV_Stock_Opportunity_Analysis': {
        'template': 'INV_Stock_Opportunity_Analysis_PRINT.docx',
        'required_fields': ['candidate.ticker', 'candidate.companyName', 'candidate.sector', 'candidate.price', 'candidate.financials', 'scorecard.weightedScore'],
        'tables': [(6, 2), (1, 1), (1, 1), (1, 1), (9, 2), (1, 1), (1, 1), (1, 1), (9, 5), (1, 1), (9, 5), (1, 1), (5, 4), (1, 1), (6, 4), (1, 1), (3, 3), (1, 1), (1, 1)],
    },
    'INV_Scenario_Analysis': {
        'template': 'INV_Scenario_Analysis_PRINT.docx',
        'required_fields': ['candidate.ticker', 'candidate.companyName', 'candidate.sector', 'candidate.price', 'scorecard.weightedScore'],
        'tables': [(6, 2), (1, 1), (7, 2), (1, 1), (9, 4), (1, 1), (5, 5), (1, 1), (2, 3), (1, 1)],
    },
    'INV_Pricing_Margin_Model': {
        'template': 'INV_Pricing_Margin_Model_PRINT.docx',
        'required_fields': ['candidate.ticker', 'candidate.companyName', 'candidate.sector', 'candidate.price', 'candidate.financials', 'candidate.revenueModel'],
        'tables': [(6, 2), (1, 1), (6, 5), (1, 1), (9, 5), (1, 1), (6, 2), (1, 1), (5, 3), (1, 1), (8, 5), (1, 1)],
    },
    'INV_Capital_Deployment': {
        'template': 'INV_Capital_Deployment_PRINT.docx',
        'required_fields': ['candidate.ticker', 'candidate.companyName', 'candidate.sector', 'candidate.price', 'scorecard.weightedScore'],
        'tables': [(6, 2), (1, 1), (8, 2), (1, 1), (7, 2), (1, 1), (5, 3), (1, 1), (5, 5), (1, 1), (6, 2), (1, 1), (1, 1)],
    },
    'INV_Rollup': {
        'template': 'INV_Rollup_PRINT.docx',
        'required_fields': ['candidates', 'candidate.ticker', 'candidate.companyName', 'candidate.sector'],
        'tables': [(1, 1), (1, 1), (1, 5), (1, 2), (1, 2), (1, 1)],
    },
}

TABLE_WIDTHS = {
    'INV_Executive_Summary': [[10800], [10800], [1550, 1550, 1550, 1550, 1550], [6000, 4800], [6000, 4800], [10800]],
    'INV_Opportunity_Scorecard': [[10800], [10800], [3200, 1200, 1200, 1200, 4200], [2500, 2500, 2500, 2500], [10800], [700, 2600, 900, 1000, 1000, 4600]],
    'INV_Stock_Opportunity_Analysis': [[10800], [10800], [10800], [10800], [2600, 8200], [10800], [2600, 2700, 2700, 2800], [10800], [3000, 1800, 1800, 1800, 2400], [10800], [2400, 3200, 1500, 3700], [10800], [3000, 3200, 1200, 3400], [10800], [3600, 3600, 3600], [2500, 2700, 2700, 2900], [10800]],
    'INV_Scenario_Analysis': [[10800], [10800], [5200, 5600], [10800], [2200, 2700, 2700, 3200], [10800], [2200, 2700, 2700, 3200], [10800], [3600, 3600, 3600]],
    'INV_Pricing_Margin_Model': [[10800], [10800], [3200, 1600, 1600, 1600, 1600], [10800], [2400, 1900, 1900, 1900, 1900], [10800], [3200, 7600], [10800], [2400, 1850, 1850, 1850, 2400], [10800], [1800, 2200, 2200, 2200, 2400]],
    'INV_Capital_Deployment': [[10800], [10800], [5200, 5600], [10800], [5200, 5600], [10800], [3200, 3600, 4000], [10800], [2400, 2400, 2400, 1800, 1800], [10800], [6800, 4000], [10800], [10800]],
    'INV_Rollup': [[10800], [10800], [10800], [700, 2400, 1000, 1000, 1000, 4700], [10800], [10800]],
}


def log_contract(report_type, template_path, tables):
    print(f'[DOCX] template={template_path}')
    for idx, (rows, cols) in enumerate(tables):
        print(f'[DOCX] {report_type} table[{idx}] rows={rows} cols={cols}')


def assert_table_shape(report_type, doc, contract):
    tables = doc.tables
    expected = contract['tables']
    if len(tables) != len(expected):
        raise SystemExit(f'{report_type} table count mismatch: expected {len(expected)}, found {len(tables)}')
    log_contract(report_type, contract['template'], expected)
    for idx, table in enumerate(tables):
        exp_rows, exp_cols = expected[idx]
        if len(table.rows) != exp_rows:
            raise SystemExit(f'{report_type} table[{idx}] row count mismatch: expected {exp_rows}, found {len(table.rows)}')
        for ridx, row in enumerate(table.rows):
            if len(row.cells) != exp_cols:
                raise SystemExit(f'{report_type} table[{idx}] row[{ridx}] column count mismatch: expected {exp_cols}, found {len(row.cells)}')


def set_table_widths(table, widths):
    if not widths:
        return
    table.autofit = False
    for row in table.rows:
        if len(row.cells) != len(widths):
            continue
        for cell, width_twips in zip(row.cells, widths):
            cell.width = Twips(int(width_twips))


def apply_table_presentation(report_type, doc):
    widths = TABLE_WIDTHS.get(report_type) or []
    for idx, table in enumerate(doc.tables):
        if idx < len(widths):
            set_table_widths(table, widths[idx])


def write_cell_lines(cell, values):
    if not isinstance(values, (list, tuple)):
        values = [values]
    lines = []
    for value in values:
        if value is None:
            continue
        text = str(value)
        parts = text.split('\n') if '\n' in text else [text]
        lines.extend(parts)
    existing = list(cell.paragraphs)
    if not existing:
        cell.add_paragraph('')
        existing = list(cell.paragraphs)
    for idx, paragraph in enumerate(existing):
        paragraph.text = lines[idx] if idx < len(lines) else ''
    for idx in range(len(existing), len(lines)):
        cell.add_paragraph(lines[idx])


def inject_cell(table, row_idx, col_idx, values, report_type, table_idx, label):
    if row_idx >= len(table.rows):
        raise SystemExit(f'{report_type} missing table[{table_idx}] row[{row_idx}] for {label}')
    row = table.rows[row_idx]
    if col_idx >= len(row.cells):
        raise SystemExit(f'{report_type} missing table[{table_idx}] row[{row_idx}] col[{col_idx}] for {label}')
    cell = row.cells[col_idx]
    write_cell_lines(cell, values)
    print(f'[DOCX] {report_type} inject {label} @ table[{table_idx}] row[{row_idx}] col[{col_idx}] <- {json.dumps(values, default=str)}')


def inject_row(table, row_idx, values, report_type, table_idx, label, start_col=0):
    if row_idx >= len(table.rows):
        raise SystemExit(f'{report_type} missing table[{table_idx}] row[{row_idx}] for {label}')
    row = table.rows[row_idx]
    if len(values) + start_col > len(row.cells):
        raise SystemExit(f'{report_type} row width mismatch for table[{table_idx}] row[{row_idx}] label={label}')
    for offset, value in enumerate(values, start=start_col):
        inject_cell(table, row_idx, offset, [value], report_type, table_idx, f'{label}[{offset}]')


def render_inv_executive_summary(doc, model):
    tables = doc.tables
    c = model['candidate'] or {}
    s = model['scorecard'] or {}
    assert_table_shape('INV_Executive_Summary', doc, CONTRACTS['INV_Executive_Summary'])

    price = float(c.get('price') or 0)
    score = float(s.get('weightedScore') or 0)
    target = round(price * (1 + max(0.12, score * 0.08)), 2) if price else None
    stop = round(price * 0.85, 2) if price else None
    rr = (target - price) / max(0.01, price - stop) if price and target and stop else None
    size_pct = float(s.get('positionSizePct') or clamp(score * 0.3, 0.5, 5.0))
    analysis_date = format_display_date(model.get('analysisDate') or c.get('analysisDate'))
    recommendation = (s.get('recommendation') or score_band(score)).upper()
    band = score_band(score)
    thesis = sentence_case(c.get('thesisSummary') or s.get('reasons', {}).get('own') or 'The current setup remains under review')
    moat = sentence_case(c.get('moatAssessment') or c.get('competitivePosition') or 'Durable differentiation is not yet fully established')
    catalysts = [c.get('catalyst1') or 'Catalyst 1 not yet identified', c.get('catalyst2') or 'Catalyst 2 not yet identified', c.get('catalyst3') or 'Catalyst 3 not yet identified']
    risks = [c.get('risk1') or 'Risk 1 not yet identified', c.get('risk2') or 'Risk 2 not yet identified', c.get('risk3') or 'Risk 3 not yet identified']
    revenue_val = c.get('financials', {}).get('revenue')
    revenue_b = float(revenue_val) / 1_000_000_000 if revenue_val is not None else None
    revenue_growth = c.get('financials', {}).get('revenueGrowth')
    gross_margin = c.get('financials', {}).get('grossMargin')
    operating_margin = c.get('financials', {}).get('operatingMargin')
    fcf_margin = c.get('financials', {}).get('fcfMargin')
    avg_target = c.get('avgTargetPrice') or c.get('targetPrice') or target
    next_earnings = format_display_date(c.get('nextEarningsDate'))
    base_downside = ((price - stop) / price) * 100 if price and stop else None
    base_upside = ((target - price) / price) * 100 if price and target else None
    bull = round(target * 1.18, 2) if target else None
    bull_upside = ((bull - price) / price) * 100 if price and bull else None

    meta_rows = [
        ('Ticker / Asset', c.get('ticker') or '—'),
        ('Sector / Asset Class', c.get('sector') or '—'),
        ('Analysis Date', analysis_date),
        ('Analyst', 'Patrick Camacho'),
        ('Current Price', money(price)),
        ('Report Version', '1.0'),
    ]
    for ridx, (label, value) in enumerate(meta_rows):
        inject_cell(tables[0], ridx, 0, [label], 'INV_Executive_Summary', 0, f'meta-label-{ridx}')
        inject_cell(tables[0], ridx, 1, [value], 'INV_Executive_Summary', 0, f'meta-value-{ridx}')

    inject_cell(tables[1], 0, 0, [
        f"{c.get('ticker') or '—'}  —  {c.get('companyName') or '—'}",
        f"{c.get('sector') or '—'}  |  Current Price: {money(price)}  |  Analysis Date: {analysis_date}",
        f"RECOMMENDATION:  {recommendation}  |  Target: {money(target)}  |  Stop: {money(stop)}  |  Size: {size_pct:.1f}%",
    ], 'INV_Executive_Summary', 1, 'summary-bar')

    cards = [
        [f'SCORECARD', f'{score:.1f} / 10', band],
        [f'BASE TARGET', money(target), f"{base_upside:.0f}% upside" if base_upside is not None else '—'],
        [f'BEAR TARGET', money(stop), f"{base_downside:.0f}% downside" if base_downside is not None else '—'],
        [f'RISK / REWARD', f'1:{rr:.1f}' if rr is not None else '—', 'Favorable' if rr is not None and rr >= 1 else 'Unfavorable'],
        [f'HORIZON', '3-6 months', 'Short-term review'],
    ]
    for idx, card in enumerate(cards):
        inject_cell(tables[2], 0, idx, card, 'INV_Executive_Summary', 2, f'card-{idx}')

    left_block = [
        'THE THESIS',
        thesis,
        'KEY CATALYSTS',
        *catalysts,
        'MOAT / EDGE',
        moat,
        f'BUY THESIS: {recommendation} at {money(price)}',
        f'RISK CHECK: {c.get("riskSummary") or "Monitor thesis-damaging developments"}',
        f'MONITORING: {c.get("monitoringApproach") or "Reassess at each catalyst window"}',
        f'NEXT CATALYST: {next_earnings or "—"}',
    ]
    right_block = [
        'KEY METRICS AT A GLANCE',
        f'Revenue (TTM):  {("$" + f"{revenue_b:.2f}B") if revenue_b is not None else "—"}',
        f'Revenue Growth:  {f"+{revenue_growth * 100:.1f}% YoY" if revenue_growth is not None else "—"}',
        f'Gross Margin:  {f"{gross_margin * 100:.1f}%" if gross_margin is not None else "—"}',
        f'Operating Margin:  {f"{operating_margin * 100:.1f}%" if operating_margin is not None else "—"}',
        f'FCF Margin:  {f"{fcf_margin * 100:.1f}%" if fcf_margin is not None else "—"}',
        f'Valuation & Consensus: P/E {num(c.get("overview", {}).get("forwardPE"), 1)}x | EV/EBITDA {num(c.get("overview", {}).get("evToEbitda"), 1)}x | Debt/Equity {num(c.get("financials", {}).get("debtToEquity"), 2)} | Range {money(c.get("low52w"))}–{money(c.get("high52w"))} | Avg PT {money(avg_target)}',
    ]
    inject_cell(tables[3], 0, 0, left_block, 'INV_Executive_Summary', 3, 'thesis-catalysts-moat')
    inject_cell(tables[3], 0, 1, right_block, 'INV_Executive_Summary', 3, 'metrics-glance')

    risk_block = [
        'TOP RISKS — THESIS BREAKERS',
        *risks,
        'WHAT TO MONITOR',
        f'Next Earnings: {format_display_date(c.get("nextEarningsDate")) if c.get("nextEarningsDate") else "—"}  |  Catalyst: {safe(c.get("catalystEvent") or c.get("catalystImpact") or "—")}',
        f'Review Cadence: {safe(c.get("monitoringApproach") or "Monthly / Quarterly")}',
    ]
    scenario_block = [
        'SCENARIO SUMMARY',
        f'Bear Case:  {money(stop)}  ({base_downside:.0f}% downside)' if base_downside is not None else 'Bear Case:  —',
        f'Base Case:  {money(target)}  (+{base_upside:.0f}% upside)' if base_upside is not None else 'Base Case:  —',
        f'Bull Case:  {money(bull)}  (+{bull_upside:.0f}% upside)' if bull_upside is not None else 'Bull Case:  —',
        'TRADE PLAN',
        f'Entry:  {money(round(price * 1.02, 2) if price else None)} — Immediate',
        f'Stop:   {money(stop)}  ({base_downside:.0f}% below entry)' if base_downside is not None else 'Stop:   —',
        f'Size:   {size_pct:.1f}% of portfolio — ${round(size_pct / 100 * 100000):,}',
        'Trim:   50% at base target / full exit at bull target or thesis break',
        f'Consensus:  {safe(c.get("analystConsensus") or s.get("consensus") or "—")}',
    ]
    inject_cell(tables[4], 0, 0, risk_block, 'INV_Executive_Summary', 4, 'risks-monitor')
    inject_cell(tables[4], 0, 1, scenario_block, 'INV_Executive_Summary', 4, 'scenario-trade')

    final_lines = [
        f'RECOMMENDATION:  {recommendation}  |  BUY AT: {money(price)}  |  STOP: {money(stop)}  |  SIZE: {size_pct:.1f}%',
        f'TARGET:  {money(target)}  |  CONSENSUS:  {safe(c.get("analystConsensus") or s.get("consensus") or "—")}',
        f'REASON:  {sentence_case(c.get("finalDecision") or c.get("decision") or f"Initiating a {size_pct:.1f}% position at {money(price)} with a {score:.1f} scorecard and {rr:.1f} risk/reward.")}',
    ]
    inject_cell(tables[5], 0, 0, final_lines, 'INV_Executive_Summary', 5, 'final-decision')

def render_inv_rollup(doc, model):

    tables = doc.tables
    c = model['candidate'] or {}
    s = model['scorecard'] or {}
    assert_table_shape('INV_Rollup', doc, CONTRACTS['INV_Rollup'])

    price = float(c.get('price') or 0)
    score = float(s.get('weightedScore') or 0)
    band = score_band(score)
    target = round(price * (1 + max(0.12, score * 0.08)), 2) if price else None
    stop = round(price * 0.85, 2) if price else None
    downside = ((price - stop) / price) * 100 if price and stop else None
    upside = ((target - price) / price) * 100 if price and target else None
    rr = (target - price) / max(0.01, price - stop) if price and target and stop else None
    size_pct = float(s.get('positionSizePct') or clamp(score * 0.3, 0.5, 5.0))
    analysis_date = format_display_date(model.get('analysisDate') or c.get('analysisDate'))
    entry_price = round(price * 1.02, 2) if price else None
    entry_display = f'${entry_price:.2f}' if entry_price else 'market'
    thesis = sentence_case(c.get('thesisSummary') or s.get('reasons', {}).get('own') or 'The current setup remains under review')
    catalysts = [
        numbered_line(1, c.get('catalyst1') or 'The first catalyst is not yet identified'),
        numbered_line(2, c.get('catalyst2') or 'The second catalyst is not yet identified'),
        numbered_line(3, c.get('catalyst3') or 'The third catalyst is not yet identified'),
    ]
    risks = [
        numbered_line(1, c.get('risk1') or 'The first risk is not yet identified'),
        numbered_line(2, c.get('risk2') or 'The second risk is not yet identified'),
        numbered_line(3, c.get('risk3') or 'The third risk is not yet identified'),
    ]
    revenue_val = c.get('financials', {}).get('revenue')
    revenue_b = float(revenue_val) / 1_000_000_000 if revenue_val is not None else None
    revenue_growth = c.get('financials', {}).get('revenueGrowth')
    gross_margin = c.get('financials', {}).get('grossMargin')
    operating_margin = c.get('financials', {}).get('operatingMargin')
    fcf_margin = c.get('financials', {}).get('fcfMargin')
    analyst_consensus = safe(c.get('analystConsensus') or s.get('analystConsensus') or s.get('consensus') or '—')
    avg_target_val = c.get('avgTargetPrice') if c.get('avgTargetPrice') is not None else c.get('targetPrice')
    if avg_target_val is None:
        avg_target_val = target
    final_decision = sentence_case(c.get('finalDecision') or c.get('decision') or f'Initiating a {size_pct:.1f}% position at {entry_display}; the {score:.1f} scorecard, 1:{rr:.1f} risk/reward, and catalyst setup justify the trade.')

    table0 = '\n'.join([
        'PERSONAL INVESTMENT RESEARCH',
        'RUP',
        'Investment Rollup',
        'One-Page Decision Brief',
        'A concise, decision-ready synthesis of all research — designed to be reviewed in under 5 minutes before executing a trade.',
        'TICKER / ASSET:  ' + (c.get('ticker') or '—'),
        'SECTOR / ASSET CLASS:  ' + (c.get('sector') or '—'),
        'ANALYSIS DATE:  ' + analysis_date,
        'ANALYST:  Patrick Camacho',
        'CURRENT PRICE:  ' + money(price),
        '',
        'NOT FINANCIAL ADVICE — FOR PERSONAL RESEARCH USE ONLY',
    ])
    table1 = (c.get('ticker') or '—') + '  —  ' + (c.get('companyName') or '—') + '\n' + (c.get('sector') or '—') + '  |  Current Price: ' + money(price) + '  |  Analysis Date: ' + analysis_date + '\nRECOMMENDATION:  ' + (s.get('recommendation') or band).upper() + '  |  Target: ' + money(target) + '  |  Stop: ' + money(stop) + '  |  Size: ' + f'{size_pct:.1f}%'
    table2_0 = 'SCORECARD\n' + f'{score:.1f} / 10' + '\n' + band
    table2_1 = 'BASE TARGET\n' + money(target) + '\n' + (f'{upside:.0f}%' if upside is not None else '—') + ' upside'
    table2_2 = 'BEAR TARGET\n' + money(stop) + '\n' + ('(' + f'{downside:.0f}%' + ') downside' if downside is not None else '—')
    table2_3 = 'RISK/REWARD\n1:' + (f'{rr:.1f}' if rr is not None else '—') + '\n' + ('Favorable' if rr is not None and rr >= 1 else 'Unfavorable')
    table2_4 = 'HORIZON\n3-6 months\nShort'
    table3_0 = 'THE THESIS\n' + thesis + '\nKEY CATALYSTS\n' + '\n'.join(catalysts) + '\nMOAT / EDGE\n' + sentence_case(c.get('moatAssessment') or c.get('keyDifferentiator') or 'Durable differentiation is not yet fully established')
    table3_1 = '\n'.join([
        'KEY METRICS AT A GLANCE',
        'Revenue (TTM):  ' + ('$' + f'{revenue_b:.2f}B' if revenue_b is not None else '—'),
        'Revenue Growth:  ' + (f'+{revenue_growth * 100:.1f}% YoY' if revenue_growth is not None else '—'),
        'Gross Margin:  ' + (f'{gross_margin * 100:.1f}%' if gross_margin is not None else '—'),
        'Operating Margin:  ' + (f'{operating_margin * 100:.1f}%' if operating_margin is not None else '—'),
        'FCF Margin:  ' + (f'{fcf_margin * 100:.1f}%' if fcf_margin is not None else '—'),
        'P/E (Forward):  ' + num(c.get('overview', {}).get('forwardPE'), 1) + 'x',
        'EV/EBITDA:  ' + num(c.get('overview', {}).get('evToEbitda'), 1) + 'x',
        'Debt/Equity:  ' + num(c.get('financials', {}).get('debtToEquity'), 2),
        '52-Week Range:  ' + money(c.get('low52w')) + ' – ' + money(c.get('high52w')),
        'Analyst Consensus:  ' + analyst_consensus + '  |  Avg PT: ' + money(avg_target_val),
    ])
    table4_0 = '\n'.join([
        'TOP RISKS — THESIS BREAKERS',
        risks[0],
        risks[1],
        risks[2],
        'WHAT TO MONITOR',
        'Next Earnings: ' + format_display_date(c.get('nextEarningsDate')) + '  |  Catalyst Event: ' + safe(c.get('catalystEvent') or c.get('catalystImpact') or '—') + '  |  Review: ' + safe(c.get('monitoringApproach') or 'Monthly / Quarterly'),
    ])
    table4_1 = '\n'.join([
        'SCENARIO SUMMARY',
        'Bear Case:  ' + money(stop) + '  (' + (f'{downside:.0f}%' if downside is not None else '—') + ')',
        'Base Case:  ' + money(target) + '  (+' + (f'{upside:.0f}%' if upside is not None else '—') + ')',
        'Bull Case:  ' + money(round(target * 1.18, 2) if target else None) + '  (+' + (f'{(((round(target * 1.18, 2) - price) / price) * 100):.0f}%' if price and target else '—') + ')',
        'TRADE PLAN',
        'Entry:  ' + money(entry_price) + ' — Immediate',
        'Stop:   ' + money(stop) + '  (' + (f'{((price - stop) / price) * 100:.0f}%' if price and stop else '—') + ' below entry)',
        'Size:   ' + f'{size_pct:.1f}% of portfolio — $' + f'{round(size_pct / 100 * 100000):,}',
        'Trim:   50% at base target / full exit at bull target or thesis break',
    ])
    table5 = 'FINAL DECISION\n' + final_decision

    inject_cell(tables[0], 0, 0, [table0], 'INV_Rollup', 0, 'masthead')
    inject_cell(tables[1], 0, 0, [table1], 'INV_Rollup', 1, 'summary-bar')
    inject_cell(tables[2], 0, 0, [table2_0], 'INV_Rollup', 2, 'scorecard')
    inject_cell(tables[2], 0, 1, [table2_1], 'INV_Rollup', 2, 'base-target')
    inject_cell(tables[2], 0, 2, [table2_2], 'INV_Rollup', 2, 'bear-target')
    inject_cell(tables[2], 0, 3, [table2_3], 'INV_Rollup', 2, 'risk-reward')
    inject_cell(tables[2], 0, 4, [table2_4], 'INV_Rollup', 2, 'horizon')
    inject_cell(tables[3], 0, 0, [table3_0], 'INV_Rollup', 3, 'thesis-catalysts-moat')
    inject_cell(tables[3], 0, 1, [table3_1], 'INV_Rollup', 3, 'metrics-glance')
    inject_cell(tables[4], 0, 0, [table4_0], 'INV_Rollup', 4, 'risks-monitor')
    inject_cell(tables[4], 0, 1, [table4_1], 'INV_Rollup', 4, 'scenario-trade')
    inject_cell(tables[5], 0, 0, [table5], 'INV_Rollup', 5, 'final-decision')


def render_inv_scorecard(doc, model):
    tables = doc.tables
    sc = model['scorecardReport']
    c = model['candidate'] or {}
    s = model['scorecard'] or {}
    assert_table_shape('INV_Opportunity_Scorecard', doc, CONTRACTS['INV_Opportunity_Scorecard'])

    analysis_date = format_display_date(model.get('analysisDate') or c.get('analysisDate'))
    price = c.get('price')
    weighted_score = s.get('weightedScore')
    final_rating = s.get('finalRating') or score_band(weighted_score)
    recommendation = s.get('recommendation') or 'PASS'
    scorecard_rows = sc.get('rows') or []
    if len(scorecard_rows) != 17:
        raise SystemExit(f'INV_Opportunity_Scorecard invalid row contract: expected 17 rows, found {len(scorecard_rows)}')

    meta_values = [c.get('ticker') or '—', c.get('sector') or '—', analysis_date, 'Patrick Camacho', money(price), '1.0']
    for ridx, value in enumerate(meta_values):
        inject_cell(tables[0], ridx, 1, [value], 'INV_Opportunity_Scorecard', 0, f'meta-{ridx}')
    inject_cell(tables[1], 0, 0, [f"SCORECARD  —  {c.get('ticker') or '—'}:  {c.get('companyName') or '—'}  ·  {money(price)}  ·  {analysis_date}"], 'INV_Opportunity_Scorecard', 1, 'title')

    for template_row in range(1, 18):
        model_row_idx = template_row - 1
        row = list(scorecard_rows[model_row_idx])
        if template_row == 16:
            row = [
                'Downside & Risk / Reward',
                '5%',
                num((s.get('rawScores') or {}).get('riskProfile'), 1),
                num((s.get('weightedScores') or {}).get('Downside / Risk-Reward'), 2),
                complete_thought(c.get('thesisBreakCondition') or c.get('riskSummary') or 'Bear case severity and risk/reward remain under review', lead='Assessment:'),
            ]
        elif template_row == 17:
            row = [
                'TOTAL WEIGHTED SCORE',
                '100%',
                '—',
                num(weighted_score, 1) + ' / 10',
                '≥8.0 STRONG BUY  ·  6–7.9 BUY  ·  4–5.9 WATCH  ·  <4.0 PASS',
            ]
        if template_row in {2, 3, 4, 6, 7, 8, 10, 11, 13, 14, 16}:
            ensure_single_line(row[0], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[0]')
            ensure_percent(row[1], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[1]')
            ensure_score_value(row[2], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[2]')
            ensure_score_value(row[3], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[3]')
            ensure_single_line(row[4], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[4]')
        elif template_row == 17:
            ensure_single_line(row[0], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[0]')
            ensure_percent(row[1], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[1]')
            ensure_single_line(row[2], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[2]')
            ensure_score_value(str(row[3]).split('/')[0], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[3]')
            ensure_single_line(row[4], f'INV_Opportunity_Scorecard table[2] row[{template_row}] col[4]')
        for cidx, value in enumerate(row):
            inject_cell(tables[2], template_row, cidx, [value], 'INV_Opportunity_Scorecard', 2, f'row-{template_row}-col-{cidx}')

    inject_cell(tables[3], 0, 0, ['STRONG BUY  8.0–10.0'], 'INV_Opportunity_Scorecard', 3, 'legend-1')
    inject_cell(tables[3], 0, 1, ['BUY  6.0–7.9'], 'INV_Opportunity_Scorecard', 3, 'legend-2')
    inject_cell(tables[3], 0, 2, ['WATCH  4.0–5.9'], 'INV_Opportunity_Scorecard', 3, 'legend-3')
    inject_cell(tables[3], 0, 3, ['PASS  < 4.0'], 'INV_Opportunity_Scorecard', 3, 'legend-4')

    inject_cell(tables[4], 0, 0, ['WATCHLIST RANKING  —  HEAD-TO-HEAD COMPARISON'], 'INV_Opportunity_Scorecard', 4, 'ranking-heading')
    inject_cell(tables[5], 0, 0, ['#'], 'INV_Opportunity_Scorecard', 5, 'rank-header-0')
    inject_cell(tables[5], 0, 1, ['TICKER / NAME'], 'INV_Opportunity_Scorecard', 5, 'rank-header-1')
    inject_cell(tables[5], 0, 2, ['SCORE'], 'INV_Opportunity_Scorecard', 5, 'rank-header-2')
    inject_cell(tables[5], 0, 3, ['BUY AT'], 'INV_Opportunity_Scorecard', 5, 'rank-header-3')
    inject_cell(tables[5], 0, 4, ['TARGET'], 'INV_Opportunity_Scorecard', 5, 'rank-header-4')
    inject_cell(tables[5], 0, 5, ['SIZE'], 'INV_Opportunity_Scorecard', 5, 'rank-header-5')
    inject_cell(tables[5], 0, 6, ['THESIS IN ONE LINE'], 'INV_Opportunity_Scorecard', 5, 'rank-header-6')
    for ridx, row in enumerate(sc.get('rankings') or []):
        if ridx + 1 >= len(tables[5].rows):
            break
        for cidx, value in enumerate(row):
            inject_cell(tables[5], ridx + 1, cidx, [value], 'INV_Opportunity_Scorecard', 5, f'rank-{ridx}-{cidx}')

    for ridx in range(1, len(tables[5].rows)):
        for cidx in range(len(tables[5].rows[ridx].cells)):
            text = tables[5].cell(ridx, cidx).text
            if '[' in text or ']' in text:
                raise SystemExit(f'INV_Opportunity_Scorecard placeholder bleed detected at table[5] row[{ridx}] col[{cidx}]: {text}')

def render_inv_stock(doc, model):
    tables = doc.tables
    stock = model['stock']
    c = model['candidate'] or {}
    s = model['scorecard'] or {}
    assert_table_shape('INV_Stock_Opportunity_Analysis', doc, CONTRACTS['INV_Stock_Opportunity_Analysis'])
    recommendation = (stock.get('conviction') or [''])[ -1].split(':', 1)[-1].strip() if stock.get('conviction') else score_band(s.get('weightedScore'))

    meta = [c.get('ticker') or '—', c.get('sector') or '—', format_display_date(model.get('analysisDate') or c.get('analysisDate')), 'Patrick Camacho', money(c.get('price')), '1.0']
    for ridx, value in enumerate(meta):
        inject_cell(tables[0], ridx, 1, [value], 'INV_Stock_Opportunity_Analysis', 0, f'meta-{ridx}')

    inject_cell(tables[1], 0, 0, ['1.  RECOMMENDATION  —  STATE THE ACTION FIRST'], 'INV_Stock_Opportunity_Analysis', 1, 'section-1')
    inject_cell(tables[2], 0, 0, [
        f"ACTION:  {stock['conviction'][-1].split(':',1)[-1].strip() if stock.get('conviction') else 'BUY / SELL / PASS'}",
        f"BUY AT:  {money(c.get('price'))}  |  TARGET:  {money(stock['targets'][1] if len(stock.get('targets') or []) > 1 else c.get('price'))}  |  STOP-LOSS:  {money(stock['targets'][0] if stock.get('targets') else c.get('price'))}",
        f"POSITION SIZE:  {stock['conviction'][1].split(':',1)[-1].strip() if len(stock.get('conviction') or []) > 1 else '—'}  |  TIME HORIZON:  {stock['conviction'][2].split(':',1)[-1].strip() if len(stock.get('conviction') or []) > 2 else '—'}",
        f"CONVICTION:  {stock['conviction'][0].split(':',1)[-1].strip() if stock.get('conviction') else 'Medium'}",
        f"RATIONALE:  {stock['final']}",
        ' ',
    ], 'INV_Stock_Opportunity_Analysis', 2, 'recommendation-block')

    inject_cell(tables[3], 0, 0, ['2.  COMPANY SNAPSHOT'], 'INV_Stock_Opportunity_Analysis', 3, 'section-2')
    snapshot_rows = [
        ('Company / Ticker', f"{c.get('companyName') or '—'} — {c.get('ticker') or '—'}"),
        ('Sector / Industry', f"{c.get('sector') or '—'} · {c.get('industry') or '—'}"),
        ('Market Cap', f"{money(c.get('marketCap'))}  |  {c.get('marketCapBucket') or '—'}"),
        ('Revenue (TTM)', f"{money(c.get('financials', {}).get('revenue'))}  |  YoY growth: {pct((c.get('financials', {}).get('revenueGrowth') or 0) * 100, 1) if c.get('financials', {}).get('revenueGrowth') is not None else '—'}"),
        ('Business Model', c.get('businessModel') or c.get('revenueModel') or '—'),
        ('Geographic Mix', c.get('geographicMix') or '—'),
        ('Key Customers', c.get('customerConcentration') or '—'),
        ('Competitive Moat', c.get('competitivePosition') or c.get('moatAssessment') or '—'),
        ('Founder / CEO', c.get('ceo') or c.get('founder') or '—'),
    ]
    for ridx, (label, value) in enumerate(snapshot_rows):
        inject_cell(tables[4], ridx, 0, [label], 'INV_Stock_Opportunity_Analysis', 4, f'snapshot-label-{ridx}')
        inject_cell(tables[4], ridx, 1, [value], 'INV_Stock_Opportunity_Analysis', 4, f'snapshot-value-{ridx}')

    inject_cell(tables[5], 0, 0, ['3.  INVESTMENT THESIS  —  WHY THIS, WHY NOW'], 'INV_Stock_Opportunity_Analysis', 5, 'section-3')
    thesis_lines = [
        f"CORE THESIS:  {stock['thesis']}",
        f"MULTI-SECTOR ANGLE:  {stock['valuationNarrative']}",
        f"CATALYST VIEW:  {stock['catalystImpact']}",
        f"MOAT VIEW:  {c.get('moatAssessment') or c.get('competitivePosition') or '—'}",
    ]
    inject_cell(tables[6], 0, 0, thesis_lines, 'INV_Stock_Opportunity_Analysis', 6, 'thesis-block')

    inject_cell(tables[7], 0, 0, ['4.  FINANCIAL PERFORMANCE'], 'INV_Stock_Opportunity_Analysis', 7, 'section-4')
    financial_rows = [
        ('Revenue', money(c.get('financials', {}).get('revenue'))),
        ('Gross Margin', pct((c.get('financials', {}).get('grossMargin') or 0) * 100, 1) if c.get('financials', {}).get('grossMargin') is not None else '—'),
        ('Operating Margin', pct((c.get('financials', {}).get('operatingMargin') or 0) * 100, 1) if c.get('financials', {}).get('operatingMargin') is not None else '—'),
        ('Net Income / EPS', money(c.get('financials', {}).get('netIncome')) if c.get('financials', {}).get('netIncome') is not None else '—'),
        ('Free Cash Flow', money(c.get('financials', {}).get('freeCashFlow')) if c.get('financials', {}).get('freeCashFlow') is not None else '—'),
        ('Debt / Equity', num(c.get('financials', {}).get('debtToEquity'), 2)),
        ('Cash & Equiv.', money(c.get('financials', {}).get('cash')) if c.get('financials', {}).get('cash') is not None else '—'),
        ('Return on Equity', pct((c.get('financials', {}).get('roe') or 0) * 100, 1) if c.get('financials', {}).get('roe') is not None else '—'),
    ]
    for ridx, (metric, value) in enumerate(financial_rows, start=1):
        inject_cell(tables[8], ridx, 0, [metric], 'INV_Stock_Opportunity_Analysis', 8, f'fin-label-{ridx}')
        inject_cell(tables[8], ridx, 1, [value], 'INV_Stock_Opportunity_Analysis', 8, f'fin-value-{ridx}-0')
        inject_cell(tables[8], ridx, 2, ['—'], 'INV_Stock_Opportunity_Analysis', 8, f'fin-value-{ridx}-1')
        inject_cell(tables[8], ridx, 3, ['—'], 'INV_Stock_Opportunity_Analysis', 8, f'fin-value-{ridx}-2')
        inject_cell(tables[8], ridx, 4, ['—'], 'INV_Stock_Opportunity_Analysis', 8, f'fin-value-{ridx}-3')

    inject_cell(tables[9], 0, 0, ['5.  VALUATION  —  IS IT CHEAP, FAIR, OR EXPENSIVE?'], 'INV_Stock_Opportunity_Analysis', 9, 'section-5')
    valuation_rows = [
        ('P/E (TTM)', num(c.get('overview', {}).get('trailingPE'), 1) + 'x', '—', '~24x', 'Fair' if c.get('overview', {}).get('trailingPE') else '—'),
        ('P/E (Forward)', num(c.get('overview', {}).get('forwardPE'), 1) + 'x', '—', '—', 'Fair' if c.get('overview', {}).get('forwardPE') else '—'),
        ('EV / EBITDA', num(c.get('overview', {}).get('evToEbitda'), 1) + 'x', '—', '—', 'Fair' if c.get('overview', {}).get('evToEbitda') else '—'),
        ('P/S Ratio', num(c.get('overview', {}).get('priceToSales'), 1) + 'x', '—', '—', 'Fair'),
        ('PEG Ratio', num(c.get('overview', {}).get('pegRatio'), 2), '—', '—', '< 1.0 = Undervalued'),
        ('P / FCF', num(c.get('overview', {}).get('priceToFreeCashFlow'), 1) + 'x', '—', '—', 'Fair'),
        ('Dividend Yield', pct((c.get('overview', {}).get('dividendYield') or 0) * 100, 2) if c.get('overview', {}).get('dividendYield') is not None else '—', '—', '~1.3%', 'Above / Below avg'),
        ('DCF Fair Value Est.', money(c.get('targetPrice') or c.get('avgTargetPrice') or c.get('price')), '—', '—', f"{((float(c.get('targetPrice') or c.get('price') or 0) - float(c.get('price') or 0)) / max(0.01, float(c.get('price') or 0)) * 100):.0f}% above/below current price"),
    ]
    for ridx, row in enumerate(valuation_rows, start=1):
        inject_cell(tables[10], ridx, 0, [row[0]], 'INV_Stock_Opportunity_Analysis', 10, f'val-label-{ridx}')
        for cidx in range(1, 5):
            inject_cell(tables[10], ridx, cidx, [row[cidx]], 'INV_Stock_Opportunity_Analysis', 10, f'val-{ridx}-{cidx}')

    inject_cell(tables[11], 0, 0, ['6.  CATALYSTS  —  WHAT MOVES THIS STOCK'], 'INV_Stock_Opportunity_Analysis', 11, 'section-6')
    catalyst_rows = [
        (stock['catalysts'][0] if len(stock.get('catalysts') or []) > 0 else '—', 'Earnings', 'Q# YYYY', stock['catalystImpact']),
        (stock['catalysts'][1] if len(stock.get('catalysts') or []) > 1 else '—', 'Ops', 'H# YYYY', stock['catalystImpact']),
        (stock['catalysts'][2] if len(stock.get('catalysts') or []) > 2 else '—', 'Macro', 'Ongoing', stock['catalystImpact']),
        (stock['catalysts'][3] if len(stock.get('catalysts') or []) > 3 else '—', 'Structural', 'YYYY+', stock['catalystImpact']),
    ]
    for ridx, row in enumerate(catalyst_rows, start=1):
        inject_cell(tables[12], ridx, 0, [row[0]], 'INV_Stock_Opportunity_Analysis', 12, f'cat-{ridx}-0')
        inject_cell(tables[12], ridx, 1, [row[1]], 'INV_Stock_Opportunity_Analysis', 12, f'cat-{ridx}-1')
        inject_cell(tables[12], ridx, 2, [row[2]], 'INV_Stock_Opportunity_Analysis', 12, f'cat-{ridx}-2')
        inject_cell(tables[12], ridx, 3, [row[3]], 'INV_Stock_Opportunity_Analysis', 12, f'cat-{ridx}-3')

    inject_cell(tables[13], 0, 0, ['7.  RISKS  —  WHAT BREAKS THE THESIS'], 'INV_Stock_Opportunity_Analysis', 13, 'section-7')
    risk_rows = [
        (stock['risks'][0] if len(stock.get('risks') or []) > 0 else '—', 'Medium', 'High', 'Sell if the first risk materializes'),
        (stock['risks'][1] if len(stock.get('risks') or []) > 1 else '—', 'Low', 'High', 'Exit if the second risk materializes'),
        (stock['risks'][2] if len(stock.get('risks') or []) > 2 else '—', 'Medium', 'Medium', 'Exit if the third risk materializes'),
        (stock['risks'][3] if len(stock.get('risks') or []) > 3 else '—', 'Low', 'Medium', 'Exit if the fourth risk materializes'),
        ('Macro / rate / FX sensitivity', 'Low', 'Medium', 'Monitor macro conditions'),
    ]
    for ridx, row in enumerate(risk_rows, start=1):
        inject_cell(tables[14], ridx, 0, [row[0]], 'INV_Stock_Opportunity_Analysis', 14, f'risk-{ridx}-0')
        inject_cell(tables[14], ridx, 1, [row[1]], 'INV_Stock_Opportunity_Analysis', 14, f'risk-{ridx}-1')
        inject_cell(tables[14], ridx, 2, [row[2]], 'INV_Stock_Opportunity_Analysis', 14, f'risk-{ridx}-2')
        inject_cell(tables[14], ridx, 3, [row[3]], 'INV_Stock_Opportunity_Analysis', 14, f'risk-{ridx}-3')

    inject_cell(tables[15], 0, 0, ['8.  PRICE TARGETS  —  BEAR / BASE / BULL'], 'INV_Stock_Opportunity_Analysis', 15, 'section-8')
    bear = money(c.get('price') * 0.85 if c.get('price') else None)
    base = money(c.get('targetPrice') or c.get('avgTargetPrice') or c.get('price'))
    bull = money((c.get('targetPrice') or c.get('price') or 0) * 1.18 if c.get('price') else None)
    inject_cell(tables[16], 0, 0, ['BEAR CASE'], 'INV_Stock_Opportunity_Analysis', 16, 'target-header-0')
    inject_cell(tables[16], 0, 1, ['BASE CASE  (PRIMARY TARGET)'], 'INV_Stock_Opportunity_Analysis', 16, 'target-header-1')
    inject_cell(tables[16], 0, 2, ['BULL CASE'], 'INV_Stock_Opportunity_Analysis', 16, 'target-header-2')
    inject_cell(tables[16], 1, 0, [f'{bear} || SELL BELOW THIS', f'({((float(c.get("price") or 0) - float(c.get("price") or 0) * 0.85) / max(0.01, float(c.get("price") or 0)) * 100):.0f}%) from current', 'Assumes downside conditions', 'Exit if thesis breaks'], 'INV_Stock_Opportunity_Analysis', 16, 'target-bear')
    inject_cell(tables[16], 1, 1, [f'{base} || TARGET — BUY AT {money(c.get("price"))}', f'+{((float(base.replace("$", "").replace(",", "")) - float(c.get("price") or 0)) / max(0.01, float(c.get("price") or 0)) * 100):.0f}% upside', 'Assumes base conditions', 'Trim discipline'], 'INV_Stock_Opportunity_Analysis', 16, 'target-base')
    inject_cell(tables[16], 1, 2, [f'{bull} || FULL UPSIDE', 'Upside extension', 'Assumes catalyst and expansion', 'Full exit discipline'], 'INV_Stock_Opportunity_Analysis', 16, 'target-bull')
    inject_cell(tables[16], 2, 0, ['Probability: 20%'], 'INV_Stock_Opportunity_Analysis', 16, 'prob-bear')
    inject_cell(tables[16], 2, 1, ['Probability: 50%'], 'INV_Stock_Opportunity_Analysis', 16, 'prob-base')
    inject_cell(tables[16], 2, 2, ['Probability: 30%'], 'INV_Stock_Opportunity_Analysis', 16, 'prob-bull')

    inject_cell(tables[17], 0, 0, ['9.  FINAL DECISION'], 'INV_Stock_Opportunity_Analysis', 17, 'section-9')
    final_lines = [
        f"RECOMMENDATION:  {recommendation}",
        ' ',
        f"BUY AT:   {money(c.get('price'))}  (current price / limit order at {money(c.get('price'))} on pullback)",
        f"SELL AT:  {money(c.get('targetPrice') or c.get('avgTargetPrice') or c.get('price'))}  (base target — trim 50% of position)",
        f"STOP:     {money(c.get('price') * 0.85 if c.get('price') else None)}  (15% below entry — hard stop, no exceptions)",
        f"SIZE:     {clamp(float(s.get('positionSizePct') or 1.0), 0.5, 5.0):.1f}% of portfolio  =  ${round(clamp(float(s.get('positionSizePct') or 1.0), 0.5, 5.0) / 100 * 100000):,}  =  [XXX] shares",
        ' ',
        f"RATIONALE:  {stock['final']}",
    ]
    inject_cell(tables[18], 0, 0, final_lines, 'INV_Stock_Opportunity_Analysis', 18, 'final-decision')
    replace_placeholders(doc, [
        f'{recommendation} — proceed' if 'BUY' in recommendation.upper() else f'{recommendation} — do not initiate',
        stock.get('valuationNarrative') or c.get('finalParagraph') or 'At current multiples, the stock trades at a fair entry point relative to its own history and growth outlook.',
        '1.0',
        money(c.get('price')),
        '100',
        '50',
    ])

def render_inv_scenario(doc, model):
    tables = doc.tables
    scenarios = model['scenarios']
    c = model['candidate'] or {}
    s = model['scorecard'] or {}
    assert_table_shape('INV_Scenario_Analysis', doc, CONTRACTS['INV_Scenario_Analysis'])

    meta = [c.get('ticker') or '—', c.get('sector') or '—', format_display_date(model.get('analysisDate') or c.get('analysisDate')), 'Patrick Camacho', money(c.get('price')), '1.0']
    for ridx, value in enumerate(meta):
        inject_cell(tables[0], ridx, 1, [value], 'INV_Scenario_Analysis', 0, f'meta-{ridx}')
    inject_cell(tables[1], 0, 0, [f"POSITION SETUP  —  {c.get('ticker') or '—'}:  {c.get('companyName') or '—'}"], 'INV_Scenario_Analysis', 1, 'title')

    setup = scenarios.get('setup') or {}
    setup_rows = [
        ('Current Market Price', f"{money(setup.get('currentPrice') or c.get('price'))}  (as of {format_display_date(model.get('analysisDate') or c.get('analysisDate'))})"),
        ('Proposed Entry Price', f"{money(setup.get('entryPrice') or c.get('price'))}  [At market / Limit order / On pullback to $XX]"),
        ('Stop-Loss', f"{money(setup.get('stopLoss') or c.get('price') * 0.85 if c.get('price') else None)}  ({((float(c.get('price') or 0) - float(setup.get('stopLoss') or c.get('price') or 0)) / max(0.01, float(c.get('price') or 0)) * 100):.0f}% below entry)  —  hard stop, no exceptions"),
        ('Shares / Units', f"{max(1, int((float(s.get('positionSizePct') or 1.0) / 100 * 100000) / max(0.01, float(c.get('price') or 1))))} shares"),
        ('Total Capital at Risk', f"{money(setup.get('capitalAtRisk') or (float(c.get('price') or 0) * 0.15))}  ({float(s.get('positionSizePct') or 1.0):.1f}% of portfolio)"),
        ('Time Horizon', f"{setup.get('timeHorizon') or '3-6 months'}  —  reassess at {format_display_date(c.get('nextEarningsDate')) if c.get('nextEarningsDate') else 'the next catalyst window'}"),
        ('Conviction Level', f"{s.get('finalRating') or score_band(s.get('weightedScore'))}  —  {c.get('finalParagraph') or c.get('thesisSummary') or 'Conviction based on current setup'}"),
    ]
    for ridx, (label, value) in enumerate(setup_rows):
        inject_cell(tables[2], ridx, 0, [label], 'INV_Scenario_Analysis', 2, f'setup-label-{ridx}')
        inject_cell(tables[2], ridx, 1, [value], 'INV_Scenario_Analysis', 2, f'setup-value-{ridx}')

    inject_cell(tables[3], 0, 0, ['THREE-SCENARIO MODEL'], 'INV_Scenario_Analysis', 3, 'section-3')
    matrix_rows = scenarios.get('tableRows') or []
    if not matrix_rows:
        matrix_rows = [
            ['Probability Weight', scenarios.get('matrix', {}).get('bear', {}).get('probability', '—'), scenarios.get('matrix', {}).get('base', {}).get('probability', '—'), scenarios.get('matrix', {}).get('bull', {}).get('probability', '—')],
            ['Revenue Growth (YoY)', scenarios.get('matrix', {}).get('bear', {}).get('growth', '—'), scenarios.get('matrix', {}).get('base', {}).get('growth', '—'), scenarios.get('matrix', {}).get('bull', {}).get('growth', '—')],
            ['Exit Multiple (P/E)', scenarios.get('matrix', {}).get('bear', {}).get('multiple', '—'), scenarios.get('matrix', {}).get('base', {}).get('multiple', '—'), scenarios.get('matrix', {}).get('bull', {}).get('multiple', '—')],
            ['EPS / EBITDA at Exit', scenarios.get('matrix', {}).get('bear', {}).get('eps', '—'), scenarios.get('matrix', {}).get('base', {}).get('eps', '—'), scenarios.get('matrix', {}).get('bull', {}).get('eps', '—')],
            ['PRICE TARGET', scenarios.get('matrix', {}).get('bear', {}).get('target', '—'), scenarios.get('matrix', {}).get('base', {}).get('target', '—'), scenarios.get('matrix', {}).get('bull', {}).get('target', '—')],
            ['Return vs. Entry', scenarios.get('matrix', {}).get('bear', {}).get('downsidePct', '—'), scenarios.get('matrix', {}).get('base', {}).get('upsidePct', '—'), scenarios.get('matrix', {}).get('bull', {}).get('upsidePct', '—')],
            ['Dollar P&L on Position', '—', '—', '—'],
            ['Key Thesis Assumption', 'What goes wrong', 'What must hold true', 'What goes right'],
        ]
    inject_cell(tables[4], 0, 0, ['SCENARIO VARIABLE'], 'INV_Scenario_Analysis', 4, 'matrix-header-0')
    inject_cell(tables[4], 0, 1, ['BEAR CASE'], 'INV_Scenario_Analysis', 4, 'matrix-header-1')
    inject_cell(tables[4], 0, 2, ['BASE CASE'], 'INV_Scenario_Analysis', 4, 'matrix-header-2')
    inject_cell(tables[4], 0, 3, ['BULL CASE'], 'INV_Scenario_Analysis', 4, 'matrix-header-3')
    for ridx, row in enumerate(matrix_rows, start=1):
        inject_cell(tables[4], ridx, 0, [row[0]], 'INV_Scenario_Analysis', 4, f'matrix-row-{ridx}-label')
        inject_cell(tables[4], ridx, 1, [row[1]], 'INV_Scenario_Analysis', 4, f'matrix-row-{ridx}-c1')
        inject_cell(tables[4], ridx, 2, [row[2]], 'INV_Scenario_Analysis', 4, f'matrix-row-{ridx}-c2')
        inject_cell(tables[4], ridx, 3, [row[3]], 'INV_Scenario_Analysis', 4, f'matrix-row-{ridx}-c3')

    inject_cell(tables[5], 0, 0, ['EXPECTED VALUE CALCULATION'], 'INV_Scenario_Analysis', 5, 'section-5')
    ev = scenarios.get('expectedValue') or {}
    ev_rows = [
        ('Bear Case', scenarios.get('matrix', {}).get('bear', {}).get('probability', '—'), scenarios.get('matrix', {}).get('bear', {}).get('target', '—'), '$[XXX]', ev.get('weightedDollar', '—')),
        ('Base Case', scenarios.get('matrix', {}).get('base', {}).get('probability', '—'), scenarios.get('matrix', {}).get('base', {}).get('target', '—'), '$[XXX]', ev.get('weightedDollar', '—')),
        ('Bull Case', scenarios.get('matrix', {}).get('bull', {}).get('probability', '—'), scenarios.get('matrix', {}).get('bull', {}).get('target', '—'), '$[XXX]', ev.get('weightedDollar', '—')),
        ('WEIGHTED EXPECTED VALUE', '100%', '—', '—', ev.get('weightedDollar', '—')),
    ]
    inject_cell(tables[6], 0, 0, ['SCENARIO'], 'INV_Scenario_Analysis', 6, 'ev-header-0')
    inject_cell(tables[6], 0, 1, ['PROBABILITY'], 'INV_Scenario_Analysis', 6, 'ev-header-1')
    inject_cell(tables[6], 0, 2, ['PRICE OUTCOME'], 'INV_Scenario_Analysis', 6, 'ev-header-2')
    inject_cell(tables[6], 0, 3, ['POSITION P&L'], 'INV_Scenario_Analysis', 6, 'ev-header-3')
    inject_cell(tables[6], 0, 4, ['EXPECTED VALUE'], 'INV_Scenario_Analysis', 6, 'ev-header-4')
    for ridx, row in enumerate(ev_rows, start=1):
        for cidx, value in enumerate(row):
            inject_cell(tables[6], ridx, cidx, [value], 'INV_Scenario_Analysis', 6, f'ev-{ridx}-{cidx}')

    inject_cell(tables[7], 0, 0, ['RISK / REWARD SUMMARY  &  DECISION'], 'INV_Scenario_Analysis', 7, 'section-7')
    bull_target = to_float(scenarios.get('matrix', {}).get('bull', {}).get('target'))
    entry_target = to_float(setup.get('entryPrice') or c.get('price'))
    stop_target = to_float(setup.get('stopLoss') or (c.get('price') * 0.85 if c.get('price') else None))
    risk_reward = [
        ('MAX DOWNSIDE', f"-{money(stop_target)}  ({((to_float(c.get('price')) or 0) - (stop_target or 0)) / max(0.01, to_float(c.get('price')) or 0) * 100:.0f}% loss from entry)"),
        ('MAX UPSIDE', f"+{money(bull_target)}  (+{scenarios.get('matrix', {}).get('bull', {}).get('upsidePct', '—')} gain)"),
        ('RISK / REWARD RATIO', f"1 : {((bull_target or 0) - (entry_target or 0)) / max(0.01, (entry_target or 0) - (stop_target or 0)):.1f}  —  {'Favorable' if float(s.get('weightedScore') or 0) >= 5.5 else 'Unfavorable'}"),
    ]
    for cidx, (label, value) in enumerate(risk_reward):
        inject_cell(tables[8], 0, cidx, [label], 'INV_Scenario_Analysis', 8, f'risk-header-{cidx}')
        inject_cell(tables[8], 1, cidx, [value], 'INV_Scenario_Analysis', 8, f'risk-value-{cidx}')

    decision = scenarios.get('decision') or {}
    decision_lines = [
        f"DECISION:  {decision.get('action', 'HOLD')}",
        ' ',
        f"{decision.get('notes') or 'State the decision clearly and tie it to the weighted EV, entry price, and risk/reward.'}",
    ]
    inject_cell(tables[9], 0, 0, decision_lines, 'INV_Scenario_Analysis', 9, 'decision-block')
    replace_placeholders(doc, [
        'At market',
        'Proceed',
        'Bull case: catalyst confirms and timing aligns',
        'Bear case: macro or execution slips',
        'Critical assumption: revenue growth and margin discipline hold',
        money(to_float(scenarios.get('matrix', {}).get('base', {}).get('target')) or to_float(c.get('price'))),
        '50',
        '15',
        '5',
        'above',
    ])

def render_inv_pmm(doc, model):
    tables = doc.tables
    c = model['candidate'] or {}
    f = c.get('financials') or {}
    assert_table_shape('INV_Pricing_Margin_Model', doc, CONTRACTS['INV_Pricing_Margin_Model'])

    meta = [c.get('ticker') or '—', c.get('sector') or '—', format_display_date(model.get('analysisDate') or c.get('analysisDate')), 'Patrick Camacho', money(c.get('price')), '1.0']
    for ridx, value in enumerate(meta):
        inject_cell(tables[0], ridx, 1, [value], 'INV_Pricing_Margin_Model', 0, f'meta-{ridx}')

    inject_cell(tables[1], 0, 0, [f'1.  REVENUE BREAKDOWN BY SEGMENT  —  {c.get("ticker") or "—"}'], 'INV_Pricing_Margin_Model', 1, 'section-1')
    revenue = float(f.get('revenue') or 0)
    rev_years = [revenue * 0.82, revenue * 0.91, revenue]
    revenue_rows = [
        ('Segment 1 — Core Product', rev_years[0] * 0.40, rev_years[1] * 0.42, rev_years[2] * 0.44, 44),
        ('Segment 2 — Recurring / SaaS / Service', rev_years[0] * 0.30, rev_years[1] * 0.30, rev_years[2] * 0.32, 32),
        ('Segment 3 — Licensing / IP', rev_years[0] * 0.18, rev_years[1] * 0.17, rev_years[2] * 0.16, 16),
        ('Segment 4 — Other / Emerging', rev_years[0] * 0.12, rev_years[1] * 0.11, rev_years[2] * 0.08, 8),
        ('TOTAL REVENUE', rev_years[0], rev_years[1], rev_years[2], 100),
    ]
    for ridx, row in enumerate(revenue_rows, start=1):
        inject_cell(tables[2], ridx, 0, [row[0]], 'INV_Pricing_Margin_Model', 2, f'revenue-label-{ridx}')
        inject_cell(tables[2], ridx, 1, [money(row[1])], 'INV_Pricing_Margin_Model', 2, f'revenue-y1-{ridx}')
        inject_cell(tables[2], ridx, 2, [money(row[2])], 'INV_Pricing_Margin_Model', 2, f'revenue-y2-{ridx}')
        inject_cell(tables[2], ridx, 3, [money(row[3])], 'INV_Pricing_Margin_Model', 2, f'revenue-y3-{ridx}')
        inject_cell(tables[2], ridx, 4, [f'{row[4]}%'], 'INV_Pricing_Margin_Model', 2, f'revenue-mix-{ridx}')

    inject_cell(tables[3], 0, 0, ['2.  MARGIN WATERFALL'], 'INV_Pricing_Margin_Model', 3, 'section-2')
    margin_rows = [
        ('Gross Margin', pct((f.get('grossMargin') or 0) * 100, 1) if f.get('grossMargin') is not None else '—', pct((f.get('grossMargin') or 0) * 100, 1) if f.get('grossMargin') is not None else '—', pct((f.get('grossMargin') or 0) * 100, 1) if f.get('grossMargin') is not None else '—', 'Expanding if mix improves'),
        ('R&D as % Revenue', '—', '—', '—', 'Investment necessary to maintain growth'),
        ('Sales & Marketing % Rev', '—', '—', '—', 'Efficiency trend depends on CAC'),
        ('G&A as % Revenue', '—', '—', '—', 'Scale leverage expected as revenue grows'),
        ('Operating Margin (EBIT)', pct((f.get('operatingMargin') or 0) * 100, 1) if f.get('operatingMargin') is not None else '—', pct((f.get('operatingMargin') or 0) * 100, 1) if f.get('operatingMargin') is not None else '—', pct((f.get('operatingMargin') or 0) * 100, 1) if f.get('operatingMargin') is not None else '—', 'Key inflection point'),
        ('EBITDA Margin', '—', '—', '—', 'Useful for capex-heavy businesses'),
        ('Net Margin', pct((f.get('netMargin') or 0) * 100, 1) if f.get('netMargin') is not None else '—', pct((f.get('netMargin') or 0) * 100, 1) if f.get('netMargin') is not None else '—', pct((f.get('netMargin') or 0) * 100, 1) if f.get('netMargin') is not None else '—', 'After-tax true profitability'),
        ('FCF Margin', pct((f.get('fcfMargin') or 0) * 100, 1) if f.get('fcfMargin') is not None else '—', pct((f.get('fcfMargin') or 0) * 100, 1) if f.get('fcfMargin') is not None else '—', pct((f.get('fcfMargin') or 0) * 100, 1) if f.get('fcfMargin') is not None else '—', 'Most important — cash conversion quality'),
    ]
    for ridx, row in enumerate(margin_rows, start=1):
        inject_cell(tables[4], ridx, 0, [row[0]], 'INV_Pricing_Margin_Model', 4, f'margin-label-{ridx}')
        for cidx in range(1, 5):
            inject_cell(tables[4], ridx, cidx, [row[cidx]], 'INV_Pricing_Margin_Model', 4, f'margin-{ridx}-{cidx}')

    inject_cell(tables[5], 0, 0, ['3.  Pricing Power Assessment'], 'INV_Pricing_Margin_Model', 5, 'section-3')
    pricing_rows = [
        ('Pricing Model', c.get('revenueModel') or 'Subscription / Volume / Value-based / Cost-plus / Dynamic'),
        ('Price Trend (3yr)', c.get('priceTrend') or 'Flat to rising'),
        ('Competitor Pricing', c.get('competitorPricing') or 'No signs of a race to the bottom'),
        ('Customer Switching Cost', c.get('switchingCost') or 'Medium'),
        ('Pricing Power Score', c.get('pricingPower') or 'Moderate'),
        ('Evidence of Power', c.get('pricingEvidence') or 'Renewals, mix shift, and steady ASPs support pricing power'),
    ]
    for ridx, row in enumerate(pricing_rows):
        inject_cell(tables[6], ridx, 0, [row[0]], 'INV_Pricing_Margin_Model', 6, f'pricing-label-{ridx}')
        inject_cell(tables[6], ridx, 1, [row[1]], 'INV_Pricing_Margin_Model', 6, f'pricing-value-{ridx}')

    inject_cell(tables[7], 0, 0, ['4.  MARGIN SENSITIVITY  —  IMPACT OF 1% CHANGE'], 'INV_Pricing_Margin_Model', 7, 'section-4')
    sens_base = float(f.get('netIncome') or f.get('income') or 0)
    sens_rows = [
        ('Gross Margin +1%', money(sens_base * 0.05), 'Driven by pricing power or mix shift'),
        ('Revenue +1%', money(sens_base * 0.03), 'Operating leverage through fixed-cost absorption'),
        ('COGS reduction 1%', money(sens_base * 0.04), 'Supply chain or automation leverage'),
        ('OpEx reduction 1%', money(sens_base * 0.03), 'Headcount or infrastructure efficiency'),
    ]
    for ridx, row in enumerate(sens_rows, start=1):
        inject_cell(tables[8], ridx, 0, [row[0]], 'INV_Pricing_Margin_Model', 8, f'sens-label-{ridx}')
        inject_cell(tables[8], ridx, 1, [row[1]], 'INV_Pricing_Margin_Model', 8, f'sens-value-{ridx}')
        inject_cell(tables[8], ridx, 2, [row[2]], 'INV_Pricing_Margin_Model', 8, f'sens-impl-{ridx}')

    inject_cell(tables[9], 0, 0, ['5.  FORWARD MARGIN PROJECTION  —  3-YEAR VIEW'], 'INV_Pricing_Margin_Model', 9, 'section-5')
    future_revenue = [revenue, revenue * 1.08, revenue * 1.16, revenue * 1.25]
    future_gm = [(f.get('grossMargin') or 0.0), (f.get('grossMargin') or 0.0) + 0.01, (f.get('grossMargin') or 0.0) + 0.02, (f.get('grossMargin') or 0.0) + 0.03]
    future_op = [(f.get('operatingMargin') or 0.0), (f.get('operatingMargin') or 0.0) + 0.01, (f.get('operatingMargin') or 0.0) + 0.02, (f.get('operatingMargin') or 0.0) + 0.03]
    future_ni = [sens_base, sens_base * 1.05, sens_base * 1.12, sens_base * 1.20]
    future_eps = [c.get('eps') or 0, (c.get('eps') or 0) * 1.05 if c.get('eps') else 0, (c.get('eps') or 0) * 1.12 if c.get('eps') else 0, (c.get('eps') or 0) * 1.20 if c.get('eps') else 0]
    future_fcf = [(f.get('fcfMargin') or 0.0), (f.get('fcfMargin') or 0.0) + 0.01, (f.get('fcfMargin') or 0.0) + 0.02, (f.get('fcfMargin') or 0.0) + 0.03]
    future_pe = [num(c.get('overview', {}).get('forwardPE'), 1), num(c.get('overview', {}).get('forwardPE'), 1), num(c.get('overview', {}).get('forwardPE'), 1), num(c.get('overview', {}).get('forwardPE'), 1)]
    projection_rows = [
        ('Revenue ($B)', [money(v) for v in future_revenue]),
        ('Gross Margin', [pct(v * 100, 1) for v in future_gm]),
        ('Operating Margin', [pct(v * 100, 1) for v in future_op]),
        ('Net Income ($M)', [money(v) for v in future_ni]),
        ('EPS', [money(v) for v in future_eps]),
        ('FCF Margin', [pct(v * 100, 1) for v in future_fcf]),
        ('Implied P/E at Entry', [f'{future_pe[0]}x', f'{future_pe[1]}x', f'{future_pe[2]}x', f'{future_pe[3]}x']),
    ]
    for ridx, row in enumerate(projection_rows, start=1):
        inject_cell(tables[10], ridx, 0, [row[0]], 'INV_Pricing_Margin_Model', 10, f'proj-label-{ridx}')
        for cidx in range(4):
            inject_cell(tables[10], ridx, cidx + 1, [row[1][cidx]], 'INV_Pricing_Margin_Model', 10, f'proj-{ridx}-{cidx}')

    inject_cell(tables[11], 0, 0, [
        'MARGIN CONCLUSION:',
        c.get('marginConclusion') or 'Margins look stable with room for expansion',
    ], 'INV_Pricing_Margin_Model', 11, 'conclusion')
    replace_placeholders(doc, [
        'Recurring revenue mix is improving as the product mix shifts toward higher-margin subscription revenue.',
        'FY25',
    ])

def render_inv_capital(doc, model):
    tables = doc.tables
    cap = model['capitalPlan']
    c = model['candidate'] or {}
    s = model['scorecard'] or {}
    assert_table_shape('INV_Capital_Deployment', doc, CONTRACTS['INV_Capital_Deployment'])
    price = to_float(c.get('price')) or 0.0
    score = float(s.get('weightedScore') or 0)
    target = round(price * (1 + max(0.12, score * 0.08)), 2) if price else None

    meta = [c.get('ticker') or '—', c.get('sector') or '—', format_display_date(model.get('analysisDate') or c.get('analysisDate')), 'Patrick Camacho', money(c.get('price')), '1.0']
    for ridx, value in enumerate(meta):
        inject_cell(tables[0], ridx, 1, [value], 'INV_Capital_Deployment', 0, f'meta-{ridx}')

    inject_cell(tables[1], 0, 0, ['1.  OPPORTUNITY SUMMARY  —  THE CASE IN BRIEF'], 'INV_Capital_Deployment', 1, 'section-1')
    base_target = to_float(c.get('targetPrice') or c.get('avgTargetPrice') or target)
    summary_rows = [
        ('Ticker / Asset', f"{c.get('ticker') or '—'} — {c.get('companyName') or '—'}"),
        ('Opportunity Scorecard', f"{num(s.get('weightedScore'), 1)} / 10  →  {score_band(s.get('weightedScore'))}"),
        ('Base Price Target', f"{money(base_target)}  (+{((base_target or 0) - (to_float(c.get('price')) or 0)) / max(0.01, to_float(c.get('price')) or 0) * 100:.0f}% upside from {money(c.get('price'))} entry)"),
        ('Stop-Loss', f"{money(c.get('price') * 0.85 if c.get('price') else None)}  (15% below entry)  —  max acceptable loss"),
        ('Risk / Reward', f"1 : {((base_target or 0) - (to_float(c.get('price')) or 0)) / max(0.01, (to_float(c.get('price')) or 0) - (to_float(c.get('price') * 0.85 if c.get('price') else None) or 0)):.1f}  →  {'Favorable' if float(s.get('weightedScore') or 0) >= 5.5 else 'Unfavorable'}"),
        ('Time Horizon', cap['summary']['horizon']),
        ('Conviction Level', f"{cap['summary']['conviction']}  —  {c.get('finalParagraph') or c.get('thesisSummary') or 'Conviction based on current setup'}"),
        ('Key Catalyst', c.get('catalystEvent') or c.get('catalyst1') or '—'),
    ]
    for ridx, (label, value) in enumerate(summary_rows):
        inject_cell(tables[2], ridx, 0, [label], 'INV_Capital_Deployment', 2, f'summary-label-{ridx}')
        inject_cell(tables[2], ridx, 1, [value], 'INV_Capital_Deployment', 2, f'summary-value-{ridx}')

    inject_cell(tables[3], 0, 0, ['2.  PORTFOLIO CONTEXT  —  BEFORE DEPLOYING'], 'INV_Capital_Deployment', 3, 'section-2')
    portfolio_rows = [
        ('Total Portfolio Value', cap['portfolio']['capital']),
        ('Available Dry Powder', cap['portfolio']['starterBudget']),
        ('Open Positions', cap['portfolio']['openPositions']),
        ('Sector Concentration', cap['portfolio']['sectorExposure']),
        ('Cash Reserve Floor', cap['portfolio']['correlationCap']),
        ('Correlated Positions', cap['portfolio']['singleNameCap']),
        ('Max Single Position Policy', cap['portfolio']['singleNameCap']),
    ]
    for ridx, (label, value) in enumerate(portfolio_rows):
        inject_cell(tables[4], ridx, 0, [label], 'INV_Capital_Deployment', 4, f'portfolio-label-{ridx}')
        inject_cell(tables[4], ridx, 1, [value], 'INV_Capital_Deployment', 4, f'portfolio-value-{ridx}')

    inject_cell(tables[5], 0, 0, ['3.  POSITION SIZE RECOMMENDATION'], 'INV_Capital_Deployment', 5, 'section-3')
    sizing_rows = [
        ('% of Portfolio', cap['sizing']['starterSize'], 'Conviction × R/R × context'),
        ('Max Dollar Risk', f"{money(float(c.get('price') or 0) * 0.15)}  (15% of portfolio at stop)", 'Never risk too much on a single name'),
        ('Shares to Buy', cap['sizing']['shares'], 'Capital ÷ entry price'),
        ('Half Kelly (Reference)', cap['sizing']['kelly'], 'Half Kelly reduces variance'),
    ]
    for ridx, row in enumerate(sizing_rows):
        inject_cell(tables[6], ridx + 1, 0, [row[0]], 'INV_Capital_Deployment', 6, f'sizing-label-{ridx}')
        inject_cell(tables[6], ridx + 1, 1, [row[1]], 'INV_Capital_Deployment', 6, f'sizing-value-{ridx}')
        inject_cell(tables[6], ridx + 1, 2, [row[2]], 'INV_Capital_Deployment', 6, f'sizing-rationale-{ridx}')

    inject_cell(tables[7], 0, 0, ['4.  STAGED ENTRY PLAN  —  DO NOT GO ALL-IN AT ONCE'], 'INV_Capital_Deployment', 7, 'section-4')
    staged_rows = [
        ('1', '50% initial', f"Buy at market / Limit at {cap['staged']['current']}", f"${round(float(c.get('price') or 0) * 0.5 * 0.1, 0):,.0f}", 'Establishes the position and lowers FOMO risk.'),
        ('2', '25% add', f"Pullback to {cap['staged']['pullback']}", f"${round(float(c.get('price') or 0) * 0.25 * 0.1, 0):,.0f}", 'Scale in on weakness.'),
        ('3', '25% add', cap['staged']['catalyst'], f"${round(float(c.get('price') or 0) * 0.25 * 0.1, 0):,.0f}", 'Only add after thesis validation.'),
        ('FULL', '100% deployed', 'All 3 tranches executed', f"${round(float(c.get('price') or 0) * 0.1, 0):,.0f}", f"{cap['staged']['maxSize']}  ·  avg cost: {money(c.get('price'))}"),
    ]
    inject_cell(tables[8], 0, 0, ['#'], 'INV_Capital_Deployment', 8, 'staged-header-0')
    inject_cell(tables[8], 0, 1, ['TRANCHE'], 'INV_Capital_Deployment', 8, 'staged-header-1')
    inject_cell(tables[8], 0, 2, ['ENTRY TRIGGER'], 'INV_Capital_Deployment', 8, 'staged-header-2')
    inject_cell(tables[8], 0, 3, ['CAPITAL'], 'INV_Capital_Deployment', 8, 'staged-header-3')
    inject_cell(tables[8], 0, 4, ['EXECUTION NOTES'], 'INV_Capital_Deployment', 8, 'staged-header-4')
    for ridx, row in enumerate(staged_rows, start=1):
        for cidx, value in enumerate(row):
            inject_cell(tables[8], ridx, cidx, [value], 'INV_Capital_Deployment', 8, f'staged-{ridx}-{cidx}')

    inject_cell(tables[9], 0, 0, ['5.  POSITION MANAGEMENT  —  RULES SET BEFORE EMOTIONS'], 'INV_Capital_Deployment', 9, 'section-5')
    rule_rows = [
        ('STOP-LOSS', cap['rules']['stop']),
        ('TRIM TARGET', cap['rules']['trim']),
        ('FULL EXIT', cap['rules']['bull']),
        ('THESIS CHECK', cap['rules']['review']),
        ('RE-ENTRY', cap['rules']['reenter']),
        ('MAX HOLD', cap['rules']['reassess']),
    ]
    for ridx, row in enumerate(rule_rows):
        inject_cell(tables[10], ridx, 0, [row[0]], 'INV_Capital_Deployment', 10, f'rule-label-{ridx}')
        inject_cell(tables[10], ridx, 1, [row[1]], 'INV_Capital_Deployment', 10, f'rule-value-{ridx}')

    inject_cell(tables[11], 0, 0, ['6.  DEPLOYMENT DECISION  —  FINAL CALL'], 'INV_Capital_Deployment', 11, 'section-6')
    decision_lines = [
        cap['decision']['action'],
        ' ',
        f"Initial deployment:  {cap['sizing']['starterSize']}  |  Stop-loss:  {cap['rules']['stop']}",
        f"Trim at:  {cap['rules']['trim']}  |  Full target:  {cap['rules']['bull']}",
        ' ',
        cap['decision']['rationale'],
    ]
    inject_cell(tables[12], 0, 0, decision_lines, 'INV_Capital_Deployment', 12, 'decision-block')

def build_model(context):

    candidate = dict(context.get('candidate') or {})
    if 'revenueModel' not in candidate and candidate.get('businessModel'):
        candidate['revenueModel'] = candidate.get('businessModel')
    if 'company_name' not in candidate and candidate.get('companyName'):
        candidate['company_name'] = candidate.get('companyName')
    if 'finalParagraph' not in candidate and candidate.get('thesisSummary'):
        candidate['finalParagraph'] = candidate.get('thesisSummary')
    scorecard = context.get('scorecard') or {}
    scenarios = make_scenarios(candidate, scorecard)
    capital_plan = make_capital_plan(candidate, scorecard)
    analysis_date = context.get('analysisDate') or ''
    candidates = context.get('candidates') or []
    return {
        'candidate': candidate,
        'candidates': candidates,
        'scorecard': scorecard,
        'scenarios': scenarios,
        'capitalPlan': capital_plan,
        'analysisDate': analysis_date,
        'executive': make_executive(candidate, scorecard),
        'scorecardReport': make_scorecard(candidate, scorecard, candidates),
        'stock': make_stock(candidate, scorecard),
        'pmm': make_pmm(candidate),
        'rollup': make_rollup(candidate, scorecard, candidates),
    }


def validate_min_size(output_path, report_type):
    size = Path(output_path).stat().st_size
    if size < MIN_DOCX_BYTES:
        raise SystemExit(f'{report_type} output too small: {size} bytes < {MIN_DOCX_BYTES}')


def validate_docx_type(output_path, report_type):
    path = Path(output_path)
    if path.suffix.lower() != '.docx':
        raise SystemExit(f'{report_type} output must end in .docx: {output_path}')
    if not zipfile.is_zipfile(path):
        raise SystemExit(f'{report_type} output is not a valid DOCX zip archive: {output_path}')
    with zipfile.ZipFile(path) as zf:
        if '[Content_Types].xml' not in zf.namelist() or 'word/document.xml' not in zf.namelist():
            raise SystemExit(f'{report_type} output is missing DOCX internals: {output_path}')


def validate_reopen(output_path, report_type):
    try:
        doc = Document(output_path)
    except Exception as exc:
        raise SystemExit(f'{report_type} reopen validation failed: {exc}') from exc
    placeholders = []
    for paragraph in iter_paragraphs(doc):
        placeholders.extend(PLACEHOLDER_RE.findall(paragraph.text or ''))
    if placeholders:
        raise SystemExit(f"{report_type} unresolved placeholders remain after save: {sorted(set(placeholders))[:12]}")


def render_docx(template_path, output_path, report_type, context):
    contract = CONTRACTS[report_type]
    model = build_model(context)
    require_fields(model, contract['required_fields'], report_type)
    doc = Document(template_path)
    assert_table_shape(report_type, doc, contract)
    if report_type == 'INV_Executive_Summary':
        render_inv_executive_summary(doc, model)
    elif report_type == 'INV_Opportunity_Scorecard':
        render_inv_scorecard(doc, model)
    elif report_type == 'INV_Stock_Opportunity_Analysis':
        render_inv_stock(doc, model)
    elif report_type == 'INV_Scenario_Analysis':
        render_inv_scenario(doc, model)
    elif report_type == 'INV_Pricing_Margin_Model':
        render_inv_pmm(doc, model)
    elif report_type == 'INV_Capital_Deployment':
        render_inv_capital(doc, model)
    elif report_type == 'INV_Rollup':
        render_inv_rollup(doc, model)
    else:
        raise SystemExit(f'Unsupported report type: {report_type}')
    apply_table_presentation(report_type, doc)
    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    validate_docx_type(out_path, report_type)
    validate_min_size(out_path, report_type)
    validate_reopen(out_path, report_type)
    print(f'[DOCX] wrote {out_path}')
    return str(out_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--template', required=True)
    parser.add_argument('--output', required=True)
    parser.add_argument('--report-type', required=True)
    parser.add_argument('--context', required=True)
    parser.add_argument('--markdown', required=True)
    args = parser.parse_args()
    context = json.loads(Path(args.context).read_text(encoding='utf-8'))
    # Markdown remains on disk for audit/debugging, but DOCX rendering now consumes structured context only.
    render_docx(args.template, args.output, args.report_type, context)
    print(args.output)


if __name__ == '__main__':
    main()
